#!/usr/bin/env python3
"""Zettelkasten VIM — Backend v4"""

import os, sys, json, base64, threading, webbrowser
import urllib.request, urllib.error, re
from pathlib import Path
from datetime import datetime
from http.server import HTTPServer, BaseHTTPRequestHandler
from socketserver import ThreadingMixIn
from urllib.parse import urlparse, unquote
import argparse

# ── Automatisch installeren van optionele PDF-pakketten ───────────────────────
def _ensure_pdf_packages():
    """Installeer pypdf, pikepdf, pdfminer.six en python-docx als ze ontbreken."""
    needed = []
    try: import pypdf
    except ImportError: needed.append("pypdf")
    try: import pikepdf
    except ImportError: needed.append("pikepdf")
    try: import pdfminer
    except ImportError: needed.append("pdfminer.six")
    try: import docx
    except ImportError: needed.append("python-docx")
    if needed:
        print(f"[setup] Installeren: {', '.join(needed)} ...", flush=True)
        import subprocess
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", "--quiet",
             "--break-system-packages"] + needed,
            capture_output=True, text=True
        )
        if result.returncode == 0:
            print(f"[setup] ✓ Geïnstalleerd: {', '.join(needed)}", flush=True)
        else:
            # Probeer zonder --break-system-packages (virtualenv)
            result2 = subprocess.run(
                [sys.executable, "-m", "pip", "install", "--quiet"] + needed,
                capture_output=True, text=True
            )
            if result2.returncode == 0:
                print(f"[setup] ✓ Geïnstalleerd (venv): {', '.join(needed)}", flush=True)
            else:
                print(f"[setup] ⚠ Installatie mislukt: {result2.stderr[:200]}", flush=True)
                print(f"[setup]   Handmatig: pip install {' '.join(needed)}", flush=True)

_ensure_pdf_packages()

DEFAULT_VAULT = Path.home() / "Zettelkasten"
STATIC_DIR    = Path(__file__).parent / "static"
VENDOR_DIR    = STATIC_DIR / "vendor"
IMAGE_EXTS    = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".svg"}

# CDN-URLs (online modus)
CDN = {
    "FONT_LINKS": (
        '<link rel="preconnect" href="https://fonts.googleapis.com">\n'
        '  <link href="https://cdn.jsdelivr.net/npm/hack-font@3/build/web/hack.css" rel="stylesheet">\n'
        '  <link href="https://fonts.googleapis.com/css2?family=DM+Sans:ital,opsz,wght@0,9..40,300;0,9..40,400;0,9..40,500;0,9..40,600;1,9..40,300&display=swap" rel="stylesheet">'
    ),
    "PDFJS_SRC":     "https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js",
    "PDFJS_WORKER":  "https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js",
    "REACT_SRC":     "https://unpkg.com/react@18/umd/react.production.min.js",
    "REACT_DOM_SRC": "https://unpkg.com/react-dom@18/umd/react-dom.production.min.js",
}

# Lokale paden (offline modus)
LOCAL = {
    "FONT_LINKS": (
        '<link href="/vendor/hack.css" rel="stylesheet">\n'
        '  <link href="/vendor/dm-sans.css" rel="stylesheet">'
    ),
    "PDFJS_SRC":     "/vendor/pdf.min.js",
    "PDFJS_WORKER":  "/vendor/pdf.worker.min.js",
    "REACT_SRC":     "/vendor/react.production.min.js",
    "REACT_DOM_SRC": "/vendor/react-dom.production.min.js",
}

def render_index(offline: bool) -> bytes:
    """Laad index.html en vul @@PLACEHOLDER@@ variabelen in op basis van modus."""
    tpl = (STATIC_DIR / "index.html").read_text(encoding="utf-8")
    values = LOCAL if offline else CDN
    for key, val in values.items():
        tpl = tpl.replace(f"@@{key}@@", val)
    return tpl.encode("utf-8")
IMAGE_MIME    = {".jpg":"image/jpeg",".jpeg":"image/jpeg",".png":"image/png",
                 ".gif":"image/gif",".webp":"image/webp",".svg":"image/svg+xml"}

class VaultManager:
    def __init__(self, vault_path):
        self.vault      = vault_path
        self.notes_dir  = vault_path / "notes"
        self.pdf_dir    = vault_path / "pdfs"
        self.annot_dir  = vault_path / "annotations"
        self.images_dir = vault_path / "images"
        self.img_annot_file = vault_path / "annotations" / "_image_annotations.json"
        self.config_file = vault_path / "config.json"
        self._init_dirs()

    def _init_dirs(self):
        for d in [self.vault, self.notes_dir, self.pdf_dir,
                  self.annot_dir, self.images_dir]:
            d.mkdir(parents=True, exist_ok=True)
        if not self.config_file.exists():
            self._write_json(self.config_file,
                {"vault_name":self.vault.name,"created":datetime.now().isoformat(),"version":4})
        if not list(self.notes_dir.glob("*.md")):
            self._create_seed()

    def _create_seed(self):
        self.save_note({"id":"20240101000001","title":"Zettelkasten — Begin hier",
            "content":"# Zettelkasten\n\n*Elke notitie is een atoom van kennis.*\n\nZie ook [[20240101000002]].\n\n#meta #start",
            "tags":["meta","start"],"created":datetime.now().isoformat(),"modified":datetime.now().isoformat()})
        self.save_note({"id":"20240101000002","title":"Links en Verbindingen",
            "content":"# Links\n\nGebruik `[[ID]]` of `[[Titel]]` voor links.\n\n#methode #links",
            "tags":["methode","links"],"created":datetime.now().isoformat(),"modified":datetime.now().isoformat()})

    # Notes
    def _note_path(self, nid): return self.notes_dir / f"{nid}.md"
    def _serialize_note(self, note):
        tags = json.dumps(note.get("tags",[]))
        fm  = f"---\nid: {note['id']}\ntitle: {note['title']}\ntags: {tags}\n"
        fm += f"created: {note.get('created',datetime.now().isoformat())}\n"
        fm += f"modified: {note.get('modified',datetime.now().isoformat())}\n"
        if note.get("sourceUrl"):   fm += f"sourceUrl: {note['sourceUrl']}\n"
        if note.get("importedAt"):  fm += f"importedAt: {note['importedAt']}\n"
        if note.get("isRead"):      fm += f"isRead: true\n"
        fm += "---\n\n"
        return fm + note.get("content","")

    def _parse_note(self, path):
        try: text = path.read_text(encoding="utf-8")
        except: return None
        note = {"id":path.stem,"title":path.stem,"tags":[],"content":"",
                "created":"","modified":"","sourceUrl":"","importedAt":"","isRead":False}
        if text.startswith("---"):
            parts = text.split("---",2)
            if len(parts)>=3:
                note["content"] = parts[2].lstrip("\n")
                for line in parts[1].strip().splitlines():
                    if   line.startswith("id:"):         note["id"]         = line[3:].strip()
                    elif line.startswith("title:"):      note["title"]      = line[6:].strip()
                    elif line.startswith("tags:"):
                        try: note["tags"] = json.loads(line[5:].strip())
                        except: note["tags"] = []
                    elif line.startswith("created:"):    note["created"]    = line[8:].strip()
                    elif line.startswith("modified:"):   note["modified"]   = line[9:].strip()
                    elif line.startswith("sourceUrl:"):  note["sourceUrl"]  = line[10:].strip()
                    elif line.startswith("importedAt:"): note["importedAt"] = line[11:].strip()
                    elif line.startswith("isRead:"):     note["isRead"]     = line[7:].strip().lower() == "true"
        else: note["content"] = text
        return note
    def load_notes(self):
        return [n for n in (self._parse_note(p) for p in
                sorted(self.notes_dir.glob("*.md"),key=lambda x:x.stat().st_mtime,reverse=True)) if n]
    def save_note(self, note):
        import tempfile as _tf, os as _os
        note["modified"] = datetime.now().isoformat()
        if not note.get("created"): note["created"] = note["modified"]
        path = self._note_path(note["id"])
        text = self._serialize_note(note)
        # Atomisch schrijven: temp + rename (veilig voor OneDrive/Google Drive sync)
        try:
            fd, tmp = _tf.mkstemp(dir=str(path.parent), suffix=".tmp")
            try:
                with _os.fdopen(fd, "w", encoding="utf-8") as f:
                    f.write(text)
                _os.replace(tmp, str(path))
            except Exception:
                try: _os.unlink(tmp)
                except: pass
                raise
        except OSError:
            path.write_text(text, encoding="utf-8")
        return note
    def delete_note(self, nid):
        p = self._note_path(nid)
        if p.exists(): p.unlink(); return True
        return False

    # Annotations
    def _annot_path(self, pdf_name):
        safe = "".join(c if c.isalnum() or c in "-_." else "_" for c in pdf_name)
        return self.annot_dir / f"{safe}.json"
    def load_annotations(self):
        r=[]
        for p in self.annot_dir.glob("*.json"):
            if p.name.startswith("_"): continue   # sla interne bestanden over
            try:
                d=json.loads(p.read_text(encoding="utf-8"))
                if isinstance(d,list): r.extend(d)
            except: pass
        return r
    def save_annotations(self, annotations):
        by_pdf={}
        for a in annotations: by_pdf.setdefault(a.get("file","unknown"),[]).append(a)
        for pdf_name,annots in by_pdf.items():
            self._annot_path(pdf_name).write_text(
                json.dumps(annots,ensure_ascii=False,indent=2),encoding="utf-8")

    def load_img_annotations(self):
        try:
            if self.img_annot_file.exists():
                d = json.loads(self.img_annot_file.read_text(encoding="utf-8"))
                return d if isinstance(d, list) else []
        except: pass
        return []

    def save_img_annotations(self, annotations):
        self.img_annot_file.write_text(
            json.dumps(annotations, ensure_ascii=False, indent=2), encoding="utf-8")

    # PDFs
    def list_pdfs(self):
        return [{"name":p.name,"size":p.stat().st_size,
                 "modified":datetime.fromtimestamp(p.stat().st_mtime).isoformat()}
                for p in sorted(self.pdf_dir.glob("*.pdf"),key=lambda x:x.stat().st_mtime,reverse=True)]
    def save_pdf(self, filename, data):
        safe="".join(c if c.isalnum() or c in "-_." else "_" for c in filename)
        (self.pdf_dir/safe).write_bytes(data); return safe
    def delete_pdf(self, filename):
        """Verwijder PDF-bestand + annotatie-bestand voor deze PDF."""
        deleted = {"pdf": False, "annotations": False}
        p = self.get_pdf_path(filename)
        if p and p.exists():
            p.unlink(); deleted["pdf"] = True
        annot_file = self._annot_path(filename)
        if annot_file.exists():
            annot_file.unlink(); deleted["annotations"] = True
        # Verwijder ook uit gecombineerde annotaties (andere PDF-bestanden)
        all_annots = self.load_annotations()
        filtered   = [a for a in all_annots if a.get("file") != filename]
        if len(filtered) != len(all_annots):
            self.save_annotations(filtered)
            deleted["annotations"] = True
        return deleted
    def get_pdf_path(self, filename):
        p=self.pdf_dir/filename; return p if p.exists() else None
    def extract_pdf_text(self, filename, max_chars=12000):
        """Extraheer tekst uit PDF — eerst puur stdlib, dan optionele packages.
        Als personal_use ingeschakeld is worden DRM-restricties genegeerd via pikepdf/qpdf."""
        p = self.get_pdf_path(filename)
        if not p: return ""

        cfg          = self.get_config()
        personal_use = cfg.get("pdf_personal_use", False)

        # ── Personal use: pikepdf of qpdf bypassen DRM-permissions ───────────
        if personal_use:
            import tempfile, os as _os

            # Methode A: pikepdf ontsleutelt → pypdf extraheert tekst
            try:
                import pikepdf, pypdf
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tf:
                    tmp_path = tf.name
                with pikepdf.open(str(p), password="",
                                  allow_overwriting_input=False) as pdf:
                    pdf.save(tmp_path)
                reader = pypdf.PdfReader(tmp_path)
                text = "\n\n".join(pg.extract_text() or "" for pg in reader.pages[:40])
                try: _os.unlink(tmp_path)
                except: pass
                print(f"[extract-A pikepdf] {len(text.strip())} tekens", flush=True)
                if len(text.strip()) > 100:
                    return text[:max_chars]
            except Exception as e:
                print(f"[extract-A pikepdf] fout: {e}", flush=True)
                try: _os.unlink(tmp_path)
                except: pass

            # Methode B: qpdf command-line → ontsleuteld PDF → pypdf
            try:
                import subprocess, pypdf
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tf:
                    tmp_path = tf.name
                r = subprocess.run(
                    ["qpdf", "--decrypt", "--stream-data=uncompress",
                     str(p), tmp_path],
                    capture_output=True, timeout=30)
                print(f"[extract-B qpdf] returncode={r.returncode} stderr={r.stderr.decode()[:100]}", flush=True)
                if r.returncode == 0:
                    reader = pypdf.PdfReader(tmp_path)
                    text = "\n\n".join(pg.extract_text() or "" for pg in reader.pages[:40])
                    try: _os.unlink(tmp_path)
                    except: pass
                    print(f"[extract-B qpdf→pypdf] {len(text.strip())} tekens", flush=True)
                    if len(text.strip()) > 100:
                        return text[:max_chars]
                try: _os.unlink(tmp_path)
                except: pass
            except Exception as e:
                print(f"[extract-B qpdf] fout: {e}", flush=True)
                try: _os.unlink(tmp_path)
                except: pass

        # ── Methode 1: pure stdlib (zlib + struct) ────────────────────────────
        try:
            text = self._extract_pdf_stdlib(p.read_bytes(), max_chars)
            print(f"[extract-1 stdlib] {len(text.strip())} tekens", flush=True)
            if len(text.strip()) > 100:
                return text
        except Exception as e:
            print(f"[extract-1 stdlib] fout: {e}", flush=True)

        # ── Methode 2: pypdf ──────────────────────────────────────────────────
        try:
            import pypdf
            reader = pypdf.PdfReader(str(p))
            if personal_use and reader.is_encrypted:
                try: reader.decrypt("")
                except Exception: pass
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages[:30])
            print(f"[extract-2 pypdf] {len(text.strip())} tekens, encrypted={reader.is_encrypted}", flush=True)
            if len(text.strip()) > 100: return text[:max_chars]
        except Exception as e:
            print(f"[extract-2 pypdf] fout: {e}", flush=True)

        # ── Methode 3: pdfminer ───────────────────────────────────────────────
        try:
            from pdfminer.high_level import extract_text as pm
            text = (pm(str(p), maxpages=20) or "")
            print(f"[extract-3 pdfminer] {len(text.strip())} tekens", flush=True)
            if len(text.strip()) > 100: return text[:max_chars]
        except Exception as e:
            print(f"[extract-3 pdfminer] fout: {e}", flush=True)

        # ── Methode 4: ruwe bytes (personal use laatste redmiddel) ────────────
        if personal_use:
            try:
                import re as _re
                raw = p.read_bytes()
                chunks = _re.findall(rb'\(([^\\\(\)]{4,200})\)', raw)
                parts = []
                for c in chunks:
                    try:
                        t = c.decode('latin-1', errors='ignore').strip()
                        if len(t) > 3 and any(ch.isalpha() for ch in t):
                            parts.append(t)
                    except: pass
                text = ' '.join(parts)
                print(f"[extract-4 rawbytes] {len(text.strip())} tekens", flush=True)
                if len(text.strip()) > 100:
                    return text[:max_chars]
            except Exception as e:
                print(f"[extract-4 rawbytes] fout: {e}", flush=True)

        print(f"[extract] ALLE methoden faalden voor {p.name}", flush=True)
        return ""

    def _extract_pdf_stdlib(self, data: bytes, max_chars=12000) -> str:
        """Verbeterde pure-Python PDF tekst extractie — geen externe packages nodig.
        Ondersteunt FlateDecode streams, Tj/TJ operatoren, ToUnicode CMap, UTF-16BE."""
        import zlib, re as _re

        def decode_str(raw):
            buf, i = [], 0
            while i < len(raw):
                if raw[i] == chr(92) and i+1 < len(raw):
                    c = raw[i+1]
                    if c.isdigit():
                        j = i+1
                        while j < i+4 and j < len(raw) and raw[j].isdigit(): j += 1
                        buf.append(chr(int(raw[i+1:j], 8) & 0xFF)); i = j; continue
                    buf.append({'n':'\n','r':'\r','t':'\t','b':'\x08','f':'\x0c',
                                '(':')',')':")",'\\\\':chr(92)}.get(c, c))
                    i += 2
                else:
                    buf.append(raw[i]); i += 1
            s = ''.join(buf)
            if len(s) >= 2 and s[0] == chr(0xFE) and s[1] == chr(0xFF):
                try: return s[2:].encode('latin-1').decode('utf-16-be', errors='replace')
                except: pass
            try:
                b = s.encode('latin-1')
                try: return b.decode('utf-8')
                except: return b.decode('latin-1', errors='replace')
            except: return s

        def inflate(raw):
            for w in (15, -15, 47):
                try: return zlib.decompress(raw, w)
                except: pass
            return raw

        def parse_cmap(txt):
            m = {}
            for blk in _re.findall(r'beginbfchar(.*?)endbfchar', txt, _re.DOTALL):
                for x in _re.finditer(r'<([0-9A-Fa-f]+)>\s*<([0-9A-Fa-f]+)>', blk):
                    try: m[int(x.group(1),16)] = chr(int(x.group(2),16))
                    except: pass
            for blk in _re.findall(r'beginbfrange(.*?)endbfrange', txt, _re.DOTALL):
                for x in _re.finditer(r'<([0-9A-Fa-f]+)>\s*<([0-9A-Fa-f]+)>\s*<([0-9A-Fa-f]+)>', blk):
                    try:
                        s,e,u = int(x.group(1),16),int(x.group(2),16),int(x.group(3),16)
                        for o in range(e-s+1): m[s+o] = chr(u+o)
                    except: pass
            return m

        def apply_cmap(s, cm):
            if not cm: return s
            out, i = [], 0
            while i < len(s):
                if i+1 < len(s):
                    c2 = (ord(s[i]) << 8) | ord(s[i+1])
                    if c2 in cm: out.append(cm[c2]); i += 2; continue
                c = ord(s[i])
                out.append(cm.get(c, s[i] if (32 <= c < 127 or c > 160) else ''))
                i += 1
            return ''.join(out)

        obj_re    = _re.compile(rb'(\d+)\s+\d+\s+obj\s*(.*?)\s*endobj', _re.DOTALL)
        stream_re = _re.compile(rb'stream\r?\n(.*?)\r?\nendstream', _re.DOTALL)
        objs = list(obj_re.finditer(data))

        # Pass 1 — CMap objecten verzamelen
        cmaps = {}
        for m in objs:
            ob = m.group(2); sm = stream_re.search(ob)
            if not sm: continue
            raw = sm.group(1)
            if b'FlateDecode' in ob[:sm.start()]: raw = inflate(raw)
            try: txt = raw.decode('latin-1', errors='replace')
            except: continue
            if 'beginbfchar' in txt or 'beginbfrange' in txt:
                cmaps[int(m.group(1))] = parse_cmap(txt)
        cmap0 = next(iter(cmaps.values())) if cmaps else {}

        # Pass 2 — tekst extraheren
        parts = []
        for m in objs:
            ob = m.group(2); sm = stream_re.search(ob)
            if not sm: continue
            raw = sm.group(1)
            if b'FlateDecode' in ob[:sm.start()]: raw = inflate(raw)
            try: dec = raw.decode('latin-1', errors='replace')
            except: continue
            for bt in _re.findall(r'BT(.*?)ET', dec, _re.DOTALL):
                bp = []
                for tj in _re.finditer(r'\(([^)]*(?:\\.[^)]*)*)\)\s*(?:Tj|\'|")', bt):
                    s = apply_cmap(decode_str(tj.group(1)), cmap0)
                    if s.strip(): bp.append(s)
                for tj in _re.finditer(r'\[(.*?)\]\s*TJ', bt, _re.DOTALL):
                    seg = []
                    for p in _re.finditer(r'\(([^)]*(?:\\.[^)]*)*)\)', tj.group(1)):
                        s = apply_cmap(decode_str(p.group(1)), cmap0)
                        if s.strip(): seg.append(s)
                        nums = _re.findall(r'(-?\d+\.?\d*)', tj.group(1)[p.end():p.end()+8])
                        if nums and float(nums[0]) < -100: seg.append(' ')
                    bp.append(''.join(seg))
                if bp: parts.append(' '.join(bp))
            parts.append('\n')
            if sum(len(x) for x in parts) > max_chars * 2: break

        result = '\n'.join(parts)
        result = _re.sub(r'[ \t]{3,}', '  ', result)
        result = _re.sub(r'\n{4,}', '\n\n\n', result)
        return result.strip()[:max_chars]

    # Images
    def list_images(self):
        imgs=[]
        for ext in IMAGE_EXTS:
            for p in self.images_dir.glob(f"*{ext}"):
                imgs.append({"name":p.name,"size":p.stat().st_size,
                    "modified":datetime.fromtimestamp(p.stat().st_mtime).isoformat(),
                    "mime":IMAGE_MIME.get(ext,"image/jpeg"),
                    "url":f"/api/image/{p.name}"})
        imgs.sort(key=lambda x:x["modified"],reverse=True); return imgs
    def save_image(self, filename, data):
        safe="".join(c if c.isalnum() or c in "-_." else "_" for c in filename)
        dest=self.images_dir/safe; dest.write_bytes(data)
        return {"name":safe,"url":f"/api/image/{safe}","size":len(data),
                "modified":datetime.now().isoformat()}
    def get_image_path(self, filename):
        p=self.images_dir/filename; return p if p.exists() else None
    def delete_image(self, filename):
        p=self.get_image_path(filename)
        if p and p.exists(): p.unlink()
        # Verwijder ook annotaties van deze afbeelding
        current = self.load_img_annotations()
        filtered = [a for a in current if a.get("file") != filename]
        if len(filtered) != len(current):
            self.save_img_annotations(filtered)
        return True
    def image_as_base64(self, filename):
        p=self.get_image_path(filename)
        if not p: return None
        ext=p.suffix.lower()
        return {"data":base64.b64encode(p.read_bytes()).decode("ascii"),
                "mime":IMAGE_MIME.get(ext,"image/jpeg")}

    # Config
    def get_config(self):
        try: return json.loads(self.config_file.read_text(encoding="utf-8"))
        except: return {}
    def save_config(self, data):
        cfg = self.get_config(); cfg.update(data)
        self._write_json(self.config_file, cfg)

    # API-sleutels — opgeslagen in config.json (hebben voorrang boven env-variabelen)
    def get_api_key(self, provider: str) -> str:
        """Geeft API-sleutel voor provider (anthropic/openai/google/openrouter).
        Volgorde: config.json → omgevingsvariabele."""
        env_map = {
            "anthropic":   "ANTHROPIC_API_KEY",
            "openai":      "OPENAI_API_KEY",
            "google":      "GOOGLE_API_KEY",
            "openrouter":  "OPENROUTER_API_KEY",
        }
        cfg_key = f"api_key_{provider}"
        val = self.get_config().get(cfg_key, "").strip()
        if val:
            return val
        return os.environ.get(env_map.get(provider, ""), "")

    def set_api_key(self, provider: str, key: str):
        self.save_config({f"api_key_{provider}": key.strip()})

    # Externe PDF-mappen (opgeslagen in config)
    def get_ext_pdf_dirs(self):
        return self.get_config().get("ext_pdf_dirs", [])
    def set_ext_pdf_dirs(self, dirs):
        self.save_config({"ext_pdf_dirs": [str(d) for d in dirs]})

    def scan_ext_pdfs(self):
        """Geeft lijst van alle PDF-bestanden in de geconfigureerde externe mappen."""
        result = []
        for d in self.get_ext_pdf_dirs():
            p = Path(d)
            if not p.is_dir(): continue
            for pdf in sorted(p.rglob("*.pdf"))[:200]:  # max 200 per map
                result.append({"name": pdf.name, "path": str(pdf),
                                "dir": str(p), "size": pdf.stat().st_size})
        return result

    def extract_pdf_text_from_path(self, abs_path: str, max_chars=8000) -> str:
        """Extraheer tekst uit een PDF op een absoluut pad (buiten de vault)."""
        p = Path(abs_path)
        if not p.exists() or not p.is_file(): return ""
        personal_use = self.get_config().get("pdf_personal_use", False)

        # Personal use: pikepdf of qpdf
        if personal_use:
            import tempfile, os as _os
            try:
                import pikepdf, pypdf
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tf:
                    tmp = tf.name
                with pikepdf.open(str(p), password="") as pdf:
                    pdf.save(tmp)
                text = "\n\n".join(pg.extract_text() or ""
                                   for pg in pypdf.PdfReader(tmp).pages[:40])
                try: _os.unlink(tmp)
                except: pass
                if len(text.strip()) > 100: return text[:max_chars]
            except Exception:
                try: _os.unlink(tmp)
                except: pass
            try:
                import subprocess, pypdf
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tf:
                    tmp = tf.name
                if subprocess.run(["qpdf","--decrypt","--stream-data=uncompress",
                                   str(p), tmp], capture_output=True, timeout=30).returncode == 0:
                    text = "\n\n".join(pg.extract_text() or ""
                                      for pg in pypdf.PdfReader(tmp).pages[:40])
                    try: _os.unlink(tmp)
                    except: pass
                    if len(text.strip()) > 100: return text[:max_chars]
                try: _os.unlink(tmp)
                except: pass
            except Exception:
                try: _os.unlink(tmp)
                except: pass

        try:
            text = self._extract_pdf_stdlib(p.read_bytes(), max_chars)
            if len(text.strip()) > 100: return text
        except Exception: pass
        try:
            import pypdf
            reader = pypdf.PdfReader(str(p))
            if personal_use and reader.is_encrypted:
                try: reader.decrypt("")
                except Exception: pass
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages[:30])
            if len(text.strip()) > 100: return text[:max_chars]
        except Exception: pass
        try:
            from pdfminer.high_level import extract_text as pm
            return (pm(str(p), maxpages=20) or "")[:max_chars]
        except Exception: pass
        if personal_use:
            try:
                import re as _re
                raw = p.read_bytes()
                chunks = _re.findall(rb'\(([^\\\(\)]{4,200})\)', raw)
                parts = []
                for c in chunks:
                    try:
                        t = c.decode('latin-1', errors='ignore').strip()
                        if len(t) > 3 and any(ch.isalpha() for ch in t):
                            parts.append(t)
                    except: pass
                text = ' '.join(parts)
                if len(text.strip()) > 100: return text[:max_chars]
            except Exception: pass
        return ""

    def extract_pdf_pages(self, filename) -> list:
        """Extraheer tekst per pagina uit vault-PDF. Geeft lijst van {page, lines, text}."""
        p = self.get_pdf_path(filename)
        if not p: return []
        data = p.read_bytes()
        return self._parse_pdf_pages(data)

    def extract_pdf_pages_from_path(self, abs_path: str) -> list:
        """Extraheer tekst per pagina uit een PDF op absoluut pad."""
        p = Path(abs_path)
        if not p.exists(): return []
        data = p.read_bytes()
        return self._parse_pdf_pages(data)

    def _parse_pdf_pages(self, data: bytes) -> list:
        """Pure-stdlib per-pagina PDF extractie. Geeft [{page, text, lines:[str]}]."""
        import zlib, re as _re

        def decode_str(raw):
            buf, i = [], 0
            while i < len(raw):
                if raw[i] == chr(92) and i+1 < len(raw):
                    c = raw[i+1]
                    if c.isdigit():
                        j = i+1
                        while j < i+4 and j < len(raw) and raw[j].isdigit(): j += 1
                        buf.append(chr(int(raw[i+1:j], 8) & 0xFF)); i = j; continue
                    buf.append({'n':'\n','r':'\r','t':'\t','b':'\x08','f':'\x0c',
                                '(':')',')':")",'\\\\':chr(92)}.get(c, c))
                    i += 2
                else:
                    buf.append(raw[i]); i += 1
            s = ''.join(buf)
            if len(s) >= 2 and s[0] == chr(0xFE) and s[1] == chr(0xFF):
                try: return s[2:].encode('latin-1').decode('utf-16-be', errors='replace')
                except: pass
            try:
                b = s.encode('latin-1')
                try: return b.decode('utf-8')
                except: return b.decode('latin-1', errors='replace')
            except: return s

        def inflate(raw):
            for w in (15, -15, 47):
                try: return zlib.decompress(raw, w)
                except: pass
            return raw

        def parse_cmap(txt):
            m = {}
            for blk in _re.findall(r'beginbfchar(.*?)endbfchar', txt, _re.DOTALL):
                for x in _re.finditer(r'<([0-9A-Fa-f]+)>\s*<([0-9A-Fa-f]+)>', blk):
                    try: m[int(x.group(1),16)] = chr(int(x.group(2),16))
                    except: pass
            for blk in _re.findall(r'beginbfrange(.*?)endbfrange', txt, _re.DOTALL):
                for x in _re.finditer(r'<([0-9A-Fa-f]+)>\s*<([0-9A-Fa-f]+)>\s*<([0-9A-Fa-f]+)>', blk):
                    try:
                        s,e,u = int(x.group(1),16),int(x.group(2),16),int(x.group(3),16)
                        for o in range(e-s+1): m[s+o] = chr(u+o)
                    except: pass
            return m

        def apply_cmap(s, cm):
            if not cm: return s
            out, i = [], 0
            while i < len(s):
                if i+1 < len(s):
                    c2 = (ord(s[i]) << 8) | ord(s[i+1])
                    if c2 in cm: out.append(cm[c2]); i += 2; continue
                c = ord(s[i])
                out.append(cm.get(c, s[i] if (32 <= c < 127 or c > 160) else ''))
                i += 1
            return ''.join(out)

        obj_re    = _re.compile(rb'(\d+)\s+\d+\s+obj\s*(.*?)\s*endobj', _re.DOTALL)
        stream_re = _re.compile(rb'stream\r?\n(.*?)\r?\nendstream', _re.DOTALL)
        page_re   = _re.compile(rb'/Type\s*/Page\b')
        objs = list(obj_re.finditer(data))

        # Pass 1 — CMap
        cmaps = {}
        for m in objs:
            ob = m.group(2); sm = stream_re.search(ob)
            if not sm: continue
            raw = sm.group(1)
            if b'FlateDecode' in ob[:sm.start()]: raw = inflate(raw)
            try: txt = raw.decode('latin-1', errors='replace')
            except: continue
            if 'beginbfchar' in txt or 'beginbfrange' in txt:
                cmaps[int(m.group(1))] = parse_cmap(txt)
        cmap0 = next(iter(cmaps.values())) if cmaps else {}

        # Pass 2 — per-pagina tekst. We groeperen content-streams op /Page objecten.
        # Simpele heuristiek: elke 'q...Q' of 'BT...ET' blok die een significante
        # tekstsprong heeft (Td/TD/Tm reset) telt als nieuwe sectie.
        # Voor pagina-toewijzing: tel /Type /Page objecten in volgorde en
        # wijs opeenvolgende content-streams toe.
        page_obj_ids = []
        for m in objs:
            ob = m.group(2)
            if page_re.search(ob[:512]):
                page_obj_ids.append(int(m.group(1)))

        # Fallback: als geen /Page objecten gevonden → alles als pagina 1
        if not page_obj_ids:
            all_text = []
            for m in objs:
                ob = m.group(2); sm = stream_re.search(ob)
                if not sm: continue
                raw = sm.group(1)
                if b'FlateDecode' in ob[:sm.start()]: raw = inflate(raw)
                try: dec = raw.decode('latin-1', errors='replace')
                except: continue
                for bt in _re.findall(r'BT(.*?)ET', dec, _re.DOTALL):
                    bp = []
                    for tj in _re.finditer(r'\(([^)]*(?:\\.[^)]*)*)\)\s*(?:Tj|\'|")', bt):
                        s2 = apply_cmap(decode_str(tj.group(1)), cmap0)
                        if s2.strip(): bp.append(s2)
                    for tj in _re.finditer(r'\[(.*?)\]\s*TJ', bt, _re.DOTALL):
                        seg = []
                        for pp in _re.finditer(r'\(([^)]*(?:\\.[^)]*)*)\)', tj.group(1)):
                            s2 = apply_cmap(decode_str(pp.group(1)), cmap0)
                            if s2.strip(): seg.append(s2)
                        bp.append(''.join(seg))
                    if bp: all_text.extend(bp)
            txt = ' '.join(all_text)
            lines = [l.strip() for l in txt.splitlines() if l.strip()]
            return [{"page":1,"text":txt,"lines":lines}]

        # Haal /Contents referenties op per Page object
        contents_re = _re.compile(rb'/Contents\s*\[([^\]]+)\]|/Contents\s+(\d+)\s+\d+\s+R')
        obj_by_id = {int(m.group(1)): m for m in objs}

        def stream_text(obj_id):
            m = obj_by_id.get(obj_id)
            if not m: return ""
            ob = m.group(2); sm = stream_re.search(ob)
            if not sm: return ""
            raw = sm.group(1)
            if b'FlateDecode' in ob[:sm.start()]: raw = inflate(raw)
            try: dec = raw.decode('latin-1', errors='replace')
            except: return ""
            parts2 = []
            for bt in _re.findall(r'BT(.*?)ET', dec, _re.DOTALL):
                bp = []
                for tj in _re.finditer(r'\(([^)]*(?:\\.[^)]*)*)\)\s*(?:Tj|\'|")', bt):
                    s2 = apply_cmap(decode_str(tj.group(1)), cmap0)
                    if s2.strip(): bp.append(s2)
                for tj in _re.finditer(r'\[(.*?)\]\s*TJ', bt, _re.DOTALL):
                    seg = []
                    for pp in _re.finditer(r'\(([^)]*(?:\\.[^)]*)*)\)', tj.group(1)):
                        s2 = apply_cmap(decode_str(pp.group(1)), cmap0)
                        if s2.strip(): seg.append(s2)
                        nums2 = _re.findall(r'(-?\d+\.?\d*)', tj.group(1)[pp.end():pp.end()+8])
                        if nums2 and float(nums2[0]) < -100: seg.append(' ')
                    bp.append(''.join(seg))
                if bp: parts2.extend(bp)
            return ' '.join(parts2)

        pages = []
        for page_num, pid in enumerate(page_obj_ids, 1):
            m = obj_by_id.get(pid)
            if not m:
                pages.append({"page": page_num, "text": "", "lines": []}); continue
            ob = m.group(2)
            cm = contents_re.search(ob)
            page_text = ""
            if cm:
                if cm.group(1):  # array
                    ids = [int(x) for x in _re.findall(rb'(\d+)\s+\d+\s+R', cm.group(1))]
                    page_text = ' '.join(stream_text(i) for i in ids)
                elif cm.group(2):
                    page_text = stream_text(int(cm.group(2)))
            if not page_text.strip():
                # Fallback: neem object erna als content
                page_text = stream_text(pid + 1)
            page_text = _re.sub(r'[ \t]{3,}', '  ', page_text)
            page_text = _re.sub(r'\n{3,}', '\n\n', page_text)
            lines = [l.strip() for l in page_text.splitlines() if l.strip()]
            pages.append({"page": page_num, "text": page_text.strip(), "lines": lines})

        return pages

    # ── PDF-tekst cache (per server-sessie) ─────────────────────────────────
    _pdf_page_cache: dict = {}

    def _cached_pdf_pages(self, fname: str) -> list:
        """extract_pdf_pages met mtime-gebaseerde cache."""
        pdf_path = self.get_pdf_path(fname)
        if not pdf_path:
            return []
        try:
            mtime = pdf_path.stat().st_mtime
        except OSError:
            return []
        key = (fname, mtime)
        if key not in VaultManager._pdf_page_cache:
            try:
                VaultManager._pdf_page_cache[key] = self.extract_pdf_pages(fname)
            except Exception:
                VaultManager._pdf_page_cache[key] = []
            # Houdt cache klein: max 30 PDFs
            if len(VaultManager._pdf_page_cache) > 30:
                oldest = next(iter(VaultManager._pdf_page_cache))
                del VaultManager._pdf_page_cache[oldest]
        return VaultManager._pdf_page_cache[key]

    def _tfidf_vectors(self, docs):
        """Berekent TF-IDF vectoren voor een lijst teksten. Geeft (vocab, matrix) terug."""
        import math, re
        stopwords = {"de","het","een","van","in","is","op","en","te","dat","die","zijn","met",
                     "voor","aan","er","maar","om","dan","als","ook","dit","nog","wel","ze",
                     "the","a","an","of","to","in","is","it","and","for","on","with","that"}
        def tokenize(t):
            return [w for w in re.findall(r"[a-z\u00c0-\u024f]{3,}", t.lower()) if w not in stopwords]
        tokenized = [tokenize(d) for d in docs]
        vocab = {w for toks in tokenized for w in toks}
        vocab = sorted(vocab)
        vidx  = {w:i for i,w in enumerate(vocab)}
        N     = len(docs)
        # IDF
        df = {}
        for toks in tokenized:
            for w in set(toks):
                df[w] = df.get(w,0)+1
        idf = {w: math.log((N+1)/(df.get(w,0)+1))+1 for w in vocab}
        # TF-IDF matrix (sparse as list of dicts)
        vecs = []
        for toks in tokenized:
            tf = {}
            for w in toks: tf[w] = tf.get(w,0)+1
            n = len(toks) or 1
            vec = {vidx[w]: (tf[w]/n)*idf[w] for w in tf if w in vidx}
            norm = math.sqrt(sum(v*v for v in vec.values())) or 1
            vecs.append({k: v/norm for k,v in vec.items()})
        return vecs

    def _cosine(self, a, b):
        return sum(a.get(k,0)*v for k,v in b.items())

    def suggest_links(self, query_content: str, current_note_id: str = "", top_n: int = 12) -> list:
        """Slim link-suggestie systeem — combineert 3 signalen:
        1. Gedeelde tags (snel, voorspelbaar)
        2. Titel-woord overlap (precies)
        3. TF-IDF cosine similariteit (semantisch)

        Geeft terug: [{id, title, tags, score, reasons:[str]}]
        Reasons leggen uit WAAROM de link relevant is.
        """
        import re as _re, math

        notes = self.load_notes()
        if not query_content.strip() or len(notes) < 2:
            return []

        q_lower = query_content.lower()

        # ── Entiteiten uit query: hoofdletterwoorden + getallen ──────────────
        entities = set(_re.findall(r'\b[A-Z][a-zA-Z]{2,}\b', query_content))
        entities |= set(_re.findall(r'\b\d{4}\b', query_content))  # jaren bijv.

        # ── TF-IDF vectoren ──────────────────────────────────────────────────
        all_texts = [n.get("title","") + " " + n.get("content","") for n in notes]
        all_texts_with_query = [query_content] + all_texts
        vecs = self._tfidf_vectors(all_texts_with_query)
        q_vec = vecs[0]
        note_vecs = vecs[1:]

        # ── Query tags (haal titels op als proxy) ────────────────────────────
        q_words = set(_re.findall(r'[a-z\u00c0-\u024f]{3,}', q_lower))

        scored = []
        for i, note in enumerate(notes):
            nid    = note.get("id","")
            title  = note.get("title","") or ""
            tags   = note.get("tags",[]) or []
            ncont  = note.get("content","") or ""

            if nid == current_note_id:
                continue

            score   = 0.0
            reasons = []

            # ── Signaal 1: TF-IDF cosine ─────────────────────────────────────
            tfidf_score = self._cosine(q_vec, note_vecs[i])
            if tfidf_score > 0.05:
                score += tfidf_score * 60
                if tfidf_score > 0.25:
                    reasons.append(f"sterk semantisch verwant ({round(tfidf_score*100)}%)")
                elif tfidf_score > 0.12:
                    reasons.append(f"semantisch verwant ({round(tfidf_score*100)}%)")

            # ── Signaal 2: Titel-woord overlap ───────────────────────────────
            title_words = set(_re.findall(r'[a-z\u00c0-\u024f]{3,}', title.lower()))
            overlap = q_words & title_words
            # Filter stopwoorden
            sw = {"de","het","een","van","in","is","op","en","te","dat","die","zijn","met",
                  "the","a","an","of","to","and","for","with","that","this","are"}
            meaningful = overlap - sw
            if meaningful:
                score += len(meaningful) * 15
                reasons.append(f"titel deelt: {', '.join(list(meaningful)[:3])}")

            # ── Signaal 3: Entiteit-overlap ───────────────────────────────────
            note_text = title + " " + ncont
            matched_entities = [e for e in entities if e in note_text]
            if matched_entities:
                score += len(matched_entities) * 20
                reasons.append(f"zelfde begrippen: {', '.join(matched_entities[:3])}")

            # ── Signaal 4: Content bevat exacte zinsdelen ────────────────────
            # Neem zinnen van 4+ woorden uit query, check of ze in notitie voorkomen
            phrases = _re.findall(r'[a-z][a-z\s]{15,40}[a-z]', q_lower)
            phrase_hits = sum(1 for p in phrases[:20] if p in ncont.lower())
            if phrase_hits:
                score += phrase_hits * 25
                reasons.append(f"deelt {phrase_hits} tekstfragment{'en' if phrase_hits>1 else ''}")

            if score > 2:
                scored.append({
                    "id":      nid,
                    "title":   title or "(geen titel)",
                    "tags":    tags,
                    "score":   round(score, 2),
                    "tfidf":   round(tfidf_score, 3),
                    "reasons": reasons[:3],  # max 3 redenen
                })

        scored.sort(key=lambda x: -x["score"])
        return scored[:top_n]


    def fuzzy_search(self, query: str, max_results=80) -> list:
        """FZF-stijl fuzzy zoeken over notities én vault-PDFs.

        Scoring (hoog = beter):
          500+ exacte phrase in titel
          300+ exacte phrase in body/paginalijn
          200   alle tokens aaneengesloten aanwezig
          100+  fuzzy karakter-voor-karakter volgorde match (bonus voor contiguïteit)
           50   tagmatch
        Resultaten: gesorteerd op score, PDF-hits gededupliceerd per (bestand, pagina).
        """
        import re as _re

        q = query.strip().lower()
        if not q:
            return []

        # Splits op spaties: elk token moet matchen (AND-logica)
        tokens = [t for t in q.split() if t]

        # ── Fuzzy-score voor één token op één regel ──────────────────────────
        def fuzzy_score_token(tok: str, line: str) -> int:
            """Geeft score >= 0 als tok fuzzy matcht in line, anders -1.
            Hogere score = betere match (exacte substring >> fuzzy spread).
            """
            if not tok:
                return 0
            # Exacte substring
            idx = line.find(tok)
            if idx >= 0:
                return 300 + max(0, 100 - idx)   # vroeg in de regel = bonus
            # Fuzzy: elk karakter van tok moet in volgorde in line voorkomen
            ti, li = 0, 0
            first_match = -1
            last_match  = -1
            consecutive = 0
            prev_match  = -1
            while ti < len(tok) and li < len(line):
                if tok[ti] == line[li]:
                    if first_match < 0:
                        first_match = li
                    if prev_match >= 0 and li == prev_match + 1:
                        consecutive += 1
                    last_match = li
                    prev_match = li
                    ti += 1
                li += 1
            if ti < len(tok):
                return -1   # niet alle chars gevonden
            spread = last_match - first_match + 1
            # Score: beloon korte spread + vroege positie + aaneengesloten
            score = max(0, 80 - spread) + max(0, 40 - first_match) + consecutive * 10
            return score if score > 0 else 1

        def score_text(text: str, toks: list) -> int:
            """Score een stuk tekst (1 regel of titel). -1 = geen match."""
            ll = text.lower()
            # Alle tokens moeten matchen
            token_scores = []
            for tok in toks:
                ts = fuzzy_score_token(tok, ll)
                if ts < 0:
                    return -1
                token_scores.append(ts)
            # Bonus voor exacte volledige phrase
            phrase = ' '.join(toks)
            phrase_bonus = 200 if phrase in ll else 0
            return sum(token_scores) + phrase_bonus

        results = []

        # ── 1. Notities ───────────────────────────────────────────────────────
        for note in self.load_notes():
            title   = note.get("title", "") or ""
            content = note.get("content", "") or ""
            tags    = note.get("tags", []) or []
            lines   = (title + "\n" + content).splitlines()

            best_score   = -1
            best_line_no = 0
            best_excerpt = ""

            # Titel krijgt extra gewicht
            title_score = score_text(title, tokens)
            if title_score >= 0:
                best_score   = title_score + 200   # titel-bonus
                best_line_no = 1
                best_excerpt = title

            # Tag-match
            tag_str = " ".join(tags)
            tag_score = score_text(tag_str, tokens) if tags else -1
            if tag_score >= 0 and tag_score + 50 > best_score:
                best_score   = tag_score + 50
                best_line_no = 1
                best_excerpt = "Tags: " + ", ".join(tags)

            # Doorzoek regels
            for ln_idx, line in enumerate(lines, 1):
                if not line.strip():
                    continue
                sc = score_text(line, tokens)
                if sc > best_score:
                    best_score   = sc
                    best_line_no = ln_idx
                    ctx_s = max(0, ln_idx - 2)
                    ctx_e = min(len(lines), ln_idx + 2)
                    best_excerpt = '\n'.join(l for l in lines[ctx_s:ctx_e] if l.strip())

            if best_score >= 0:
                results.append({
                    "type":    "note",
                    "id":      note.get("id"),
                    "title":   title or "(geen titel)",
                    "tags":    tags,
                    "score":   best_score,
                    "line":    best_line_no,
                    "excerpt": best_excerpt[:300],
                    "content": content,
                })

        # ── 2. Vault PDFs (met cache) ─────────────────────────────────────────
        for pdf_info in self.list_pdfs():
            fname = pdf_info["name"]
            pages = self._cached_pdf_pages(fname)

            # Beste hit per (bestand, pagina) voor deduplicatie
            best_per_page: dict = {}

            for pg in pages:
                page_num = pg["page"]
                pg_lines = pg.get("lines", [])

                for ln_idx, line in enumerate(pg_lines, 1):
                    if not line.strip():
                        continue
                    sc = score_text(line, tokens)
                    if sc < 0:
                        continue
                    pg_key = (fname, page_num)
                    if pg_key not in best_per_page or sc > best_per_page[pg_key]["score"]:
                        ctx_s   = max(0, ln_idx - 2)
                        ctx_e   = min(len(pg_lines), ln_idx + 2)
                        excerpt = '\n'.join(
                            l for l in pg_lines[ctx_s:ctx_e] if l.strip()
                        )
                        best_per_page[pg_key] = {
                            "type":    "pdf",
                            "source":  fname,
                            "score":   sc,
                            "page":    page_num,
                            "line":    ln_idx,
                            "excerpt": excerpt[:300],
                            "title":   f"{fname}  —  p.{page_num}",
                        }

            results.extend(best_per_page.values())

        results.sort(key=lambda x: -x["score"])
        return results[:max_results]

    # ── Spell engine (singleton, lazy-loaded per taal) ──────────────────────────
    _spell_cache: dict = {}     # {"en": set(...), "nl": set(...)}
    _spell_lock = None

    def fulltext_search(self, query: str, max_results=50) -> list:
        """Full-text zoeken: alle overeenkomende regels per notitie,
        met geselecteerde context en gemarkeerde hits."""
        import re as _re

        q = query.strip()
        if not q:
            return []

        # Bouw regex — escape special chars, case-insensitive
        try:
            pattern = _re.compile(_re.escape(q), _re.IGNORECASE)
        except Exception:
            return []

        results = []

        for note in self.load_notes():
            title   = note.get("title", "") or ""
            content = note.get("content", "") or ""
            tags    = note.get("tags", []) or []
            lines   = content.splitlines()

            matches = []  # [{line_no, line, context}]

            for ln_idx, line in enumerate(lines):
                if pattern.search(line):
                    # Context: 1 regel voor en na
                    ctx_start = max(0, ln_idx - 1)
                    ctx_end   = min(len(lines), ln_idx + 2)
                    context   = "\n".join(
                        l for l in lines[ctx_start:ctx_end] if l.strip()
                    )
                    matches.append({
                        "line_no":  ln_idx + 1,
                        "line":     line.strip()[:200],
                        "context":  context[:400],
                    })

            # Titel match telt ook
            title_match = bool(pattern.search(title))

            if matches or title_match:
                # Score: titel-match hoog, dan op aantal matches
                score = (500 if title_match else 0) + len(matches) * 10
                results.append({
                    "type":       "note",
                    "id":         note.get("id"),
                    "title":      title or "(geen titel)",
                    "tags":       tags,
                    "score":      score,
                    "matches":    matches[:20],  # max 20 hits per notitie
                    "match_count": len(matches) + (1 if title_match else 0),
                    "title_match": title_match,
                    "content":    content,
                })

        results.sort(key=lambda r: r["score"], reverse=True)
        return results[:max_results]


    @classmethod
    def _get_lock(cls):
        import threading
        if cls._spell_lock is None:
            cls._spell_lock = threading.Lock()
        return cls._spell_lock

    @classmethod
    def _load_hunspell(cls, dic_path: str, aff_path: str) -> set:
        """Laad een Hunspell .dic/.aff bestand en genereer alle woordvormen via suffix-regels."""
        import re as _re
        sfx = {}
        try:
            with open(aff_path, encoding="utf-8", errors="replace") as f:
                lines = f.readlines()
            i = 0
            while i < len(lines):
                m = _re.match(r"^(SFX)\s+(\S+)\s+[YN]\s+(\d+)", lines[i].strip())
                if m:
                    flag, n = m.group(2), int(m.group(3))
                    rules = []
                    for j in range(1, n + 1):
                        if i + j < len(lines):
                            rm = _re.match(r"^SFX\s+\S+\s+(\S+)\s+(\S+)\s+(\S+)",
                                           lines[i + j].strip())
                            if rm:
                                s, a, c = rm.group(1), rm.group(2), rm.group(3)
                                rules.append(("" if s == "0" else s,
                                              "" if a == "0" else a, c))
                    sfx[flag] = rules
                    i += n + 1; continue
                i += 1
        except Exception:
            pass
        words = set()
        try:
            with open(dic_path, encoding="utf-8", errors="replace") as f:
                next(f, None)
                for line in f:
                    line = line.strip()
                    if not line: continue
                    word  = line.split("/")[0].lower()
                    flags = line.split("/")[1] if "/" in line else ""
                    words.add(word)
                    for flag in flags:
                        for s, a, c in sfx.get(flag, []):
                            try:
                                check = word[:-len(s)] if s and word.endswith(s) else word
                                if not check: continue
                                if c == "." or _re.search(c + "$", check):
                                    words.add(check + a)
                            except Exception:
                                pass
        except Exception:
            pass
        return words

    @classmethod
    def _nl_builtin_words(cls) -> set:
        """Ingebakken Nederlandse woordenlijst — ~3000 veelgebruikte woorden.
        Bevat expliciete vormen (geen algoritme-expansie die nonsens genereert).
        Fallback als er geen Hunspell .dic beschikbaar is.
        """
        # Meest voorkomende Nederlandse woorden + werkwoorden + bijvoeglijke naamwoorden
        # Elke vorm is expliciet opgenomen — geen morfologie-expansie
        words = """
aan aanbieden aanbod aandacht aandeel aanpak aanpassen aanwezig aard achter
achtergrond actie actueel adres afbeelding afdeling afgelopen afhangen afmaken
afstand afstemmen agenda algoritme algemeen alleen ander anders antwoord argument
aspect automatisch auto avond baan bedoeling beeld begrijpen belang belangrijk
beleid bereiken beschikbaar beschrijven bespreken bestand bestuur betrekking
bewust bijdrage bijlage bijzonder boek bron buiten bureau categorie code
combinatie commercieel communicatie component conclusie conditie configuratie
context contract controle dagelijks data datum definitie detail digitaal doel
document doorlopend duidelijk effect effectief efficiënt element evaluatie
factor fout framework functie gebruik gebruiken gedachte gelijk gevaar gegevens
gezond goed groep groot hardware historisch hoofd hypothese idee implementatie
informatie integratie intern inzicht jaar keuze kiezen koppeling kwaliteit
later leven link maand manier meer mensen methode middel minder model module
moment naam netwerk niveau normaal object observatie omgeving ondersteuning
ontwerp ontwikkelen oplossing organisatie overzicht pagina parameter periode
perspectief plan positief prioriteit principe probleem proces programma protocol
punt realiseren resultaat richting rol samenwerking schema scope sectie server
situatie sleutel soort specifiek sport start structuur systeem taak technisch
tekst termijn theorie tijd toelichting toepassing type uitvoeren uiteindelijk
validatie versie verwachting voorstel voorbeeld vraag waarde werkwijze
de het een en van in is dat dit met op aan voor uit door ook naar toe zo nu
ja nee ik je jij hij zij ze we hen hun hem zijn haar ons onze jullie u
bij door in met na naar om op over te tegen tot voor via zonder
achter beneden binnen boven buiten langs naast
die welke waar wanneer hoe wie wat waarom waardoor waarvoor
dan als of want omdat zodat tenzij totdat voordat nadat hoewel
maar toch ook wel nog altijd nooit soms even zeker misschien echter bovendien
al erg heel zeer vrij nogal bijna reeds toen thans steeds dus
gaan gaat gingen gegaan komen komt kwamen worden wordt werden hebben heeft
kunnen kan konden moeten moet moesten mogen mag mochten willen wil wilden
zullen zal zouden laten laat lieten zien ziet zagen doen doet deden
nemen neemt namen geven geeft gaven staan staat stonden liggen ligt lagen
lopen loopt liepen werken werkt werkten schrijven schrijft schreven
lezen leest lazen spreken spreekt spraken beginnen begint begonnen
stoppen stopt stopten vragen vraagt vroegen antwoorden helpen helpt hielpen
zoeken zoekt zochten vinden vindt vonden kennen kent kenden weten weet wisten
hopen hoopt hoopten verwachten probeert proberen gebruiken maakt bouwen bouwt
groot groter grootst klein kleiner kleinst goed beter best slecht slechter
nieuw nieuwe oud oude eerste tweede laatste volgende vorige huidige
hoog hoge laag lage lang lange kort korte breed brede dik dikke
zwaar zware licht lichte snel snelle langzaam langzame mooi mooie
sterk sterke hard harde zacht zachte warm warme koud koude rustig stille
vrolijk vrolijke triest droevige gelukkig ongelukkige
academisch administratief beschikbaar complex digitaal effectief
extern functioneel historisch informatief intern logisch operationeel
organisatorisch technisch transparant verantwoordelijk wetenschappelijk
aanbeveling aanvraag aanpassing afspraak afhandeling berekening beschrijving
bijwerking boodschap controle correctie discussie doelstelling foutmelding
gebruik handleiding herziening instructie klacht melding ontwerp opmerking
overleg planning presentatie procedure rapport terugkoppeling uitleg
verbetering verslag voortgang wijziging zoekopdracht
huis huizen fiets fietsen auto autos school scholen werk werken
man mannen vrouw vrouwen kind kinderen dag dagen week weken
water vuur aarde lucht licht donker wit zwart rood blauw groen geel
stad steden land landen weg wegen tijd tijden naam namen
getal getallen woord woorden zin zinnen tekst teksten
nieuw oud groot klein snel langzaam mooi lelijk goed slecht
echt nep zeker onzeker mogelijk onmogelijk nodig nuttig handig
altijd nooit soms vaak zelden bijna precies ongeveer exact
hier daar ergens nergens overal tegelijk samen apart
iemand niemand iedereen alles niets veel weinig genoeg meer minder
door via met zonder voor na tijdens voor gedurende
beste slechtste mooiste lelijkste grootste kleinste
worden geweest gedaan gezegd gegeven genomen gelaten gekomen
gemaakt gebouwd gevonden gezocht gevraagd gewerkt gelopen
schrijven lezen spreken luisteren kijken zien horen voelen
denken weten begrijpen leren onderwijzen uitleggen beschrijven
starten stoppen beginnen eindigen doorgaan hervatten afbreken
openen sluiten opslaan laden importeren exporteren downloaden
toevoegen verwijderen wijzigen aanpassen kopiëren plakken
zoeken vinden filteren sorteren ordenen rangschikken
verbinden koppelen integreren synchroniseren updaten
berekenen analyseren evalueren testen valideren controleren
plannen organiseren coördineren beheren onderhouden
communiceren samenwerken overleggen afstemmen rapporteren
"""
        return set(w.strip() for w in words.split() if len(w.strip()) >= 2)

    @classmethod
    def _ensure_spell(cls, lang: str, vault_dir=None):
        """Laad en cache woordenlijst voor taal (eenmalig per server-run).

        Zoekpaden per taal — in volgorde van prioriteit:
          0. static/vendor/dict/   — ingebakken in de app (via download-dictionaries.sh)
          1. vault_dir/<lang>.dic  — eigen woordenlijst in vault
          2. /opt/homebrew/share/hunspell/
          3. LibreOffice bundled   — /Applications/LibreOffice.app/...
          4. /usr/share/hunspell/  — Linux systeem
          5. /usr/share/myspell/   — oudere Linux distro's
          6. Ingebakken fallback   — _nl_builtin_words() (alleen NL)
        """
        with cls._get_lock():
            if lang in cls._spell_cache:
                return
            from pathlib import Path as _P
            import glob as _glob

            # ── Pad naar static/vendor/dict/ (naast server.py) ──────────────
            _server_dir = _P(__file__).parent
            _vendor_dict = _server_dir / "static" / "vendor" / "dict"

            words = set()
            if lang == "en":
                candidates = [
                    str(_vendor_dict / "en_US.dic"),          # 0. ingebakken in app
                    "/opt/homebrew/share/hunspell/en_US.dic", # 2. Homebrew
                    "/usr/share/hunspell/en_US.dic",          # 4. Linux
                    "/usr/share/myspell/en_US.dic",           # 5. oudere Linux
                ]
                for lo in _glob.glob("/Applications/LibreOffice*.app/Contents/Resources/extensions/dict-en/*/en_US.dic"):
                    candidates.insert(2, lo)                  # 3. LibreOffice Mac
                if vault_dir:
                    candidates.insert(1, str(_P(vault_dir) / "en_US.dic"))  # 1. vault
                for dic in candidates:
                    aff = dic.replace(".dic", ".aff")
                    if _P(dic).exists() and _P(aff).exists():
                        words = cls._load_hunspell(dic, aff)
                        print(f"[spell] EN geladen: {dic} ({len(words)} woorden)", flush=True)
                        break

            elif lang == "nl":
                candidates = [
                    str(_vendor_dict / "nl_NL.dic"),          # 0. ingebakken in app
                    "/opt/homebrew/share/hunspell/nl_NL.dic", # 2. Homebrew
                    str(_P.home() / "Library/Spelling/nl_NL.dic"),  # macOS user
                    "/Library/Spelling/nl_NL.dic",            # macOS systeem
                    "/usr/share/hunspell/nl_NL.dic",          # 4. Linux
                    "/usr/share/myspell/nl_NL.dic",           # 5. oudere Linux
                ]
                for lo in _glob.glob("/Applications/LibreOffice*.app/Contents/Resources/extensions/dict-nl/*/nl_NL.dic"):
                    candidates.insert(2, lo)                  # 3. LibreOffice Mac
                if vault_dir:
                    candidates.insert(1, str(_P(vault_dir) / "nl_NL.dic"))  # 1. vault
                loaded = False
                for dic in candidates:
                    aff = dic.replace(".dic", ".aff")
                    if _P(dic).exists() and _P(aff).exists():
                        words = cls._load_hunspell(dic, aff)
                        print(f"[spell] NL geladen: {dic} ({len(words)} woorden)", flush=True)
                        loaded = True
                        break
                if not loaded:
                    words = cls._nl_builtin_words()
                    print(
                        f"[spell] NL: geen woordenboek gevonden — gebruik ingebakken fallback "
                        f"({len(words)} woorden).\n"
                        f"         Tip: voer uit: cd static/vendor && bash download-dictionaries.sh",
                        flush=True
                    )
            cls._spell_cache[lang] = words

    @classmethod
    def _suggest(cls, word: str, wordset: set, max_n: int = 8) -> list:
        """Genereer spellingssuggesties via edit-distance 1 en 2."""
        import re as _re
        alphabet = "abcdefghijklmnopqrstuvwxyz"
        def edits1(w):
            splits  = [(w[:i], w[i:]) for i in range(len(w)+1)]
            deletes = [L+R[1:]       for L,R in splits if R]
            transposes=[L+R[1]+R[0]+R[2:] for L,R in splits if len(R)>1]
            replaces= [L+c+R[1:]    for L,R in splits if R for c in alphabet]
            inserts = [L+c+R        for L,R in splits       for c in alphabet]
            return set(deletes+transposes+replaces+inserts)
        candidates = edits1(word) & wordset
        if len(candidates) < max_n:
            for e1 in edits1(word):
                candidates |= edits1(e1) & wordset
        # Sorteer op overeenkomst met origineel (langere gemeenschappelijke prefix eerst)
        def score(c):
            common = sum(1 for a,b in zip(word,c) if a==b)
            return (-common, abs(len(c)-len(word)))
        return sorted(candidates, key=score)[:max_n]

    def spellcheck_words(self, words: list, lang: str = "en") -> dict:
        """Spellcheck via geladen Hunspell-woordenboek.
        Resultaat: {word: {"spell": bool, "suggestions": [str]}}
        
        Slimmer dan voorheen:
        - Hoofdletterwoorden worden ook gecheckt (lowercase variant)
        - Samengestelde woorden (NL: "woordenboek") splitst op deelwoorden
        - Fonetisch vergelijkbare suggesties bovenaan
        - Resultaat geïndexeerd op ZOWEL originele als lowercase vorm
        """
        import re as _re
        VaultManager._ensure_spell(lang, vault_dir=str(self.vault))
        known = VaultManager._spell_cache.get(lang, set())

        # Vault-woorden zijn altijd correct (eigen termen, namen, afkortingen)
        vault_ok: set = set()
        for note in self.load_notes():
            text = (note.get("title","") + " " + note.get("content","")).lower()
            for w in _re.findall(r"[a-zA-ZÀ-ɏ'\-]{2,}", text):
                vault_ok.add(w.strip("'-"))

        always_ok = _re.compile(
            r"^\d"
            r"|^[A-Z]{2,}$"          # afkortingen (NATO, PDF)
            r"|^[a-z]{1,2}$"         # de, en, in, op
            r"|https?://"
            r"|\[\["
            r"|^@\w"
            r"|^\d+[a-z]+$"         # 3d, 2e, 10px
        )

        def is_compound_ok(w, known_set):
            """NL samengestelde woorden: 'woordenboek' → 'woorden'+'boek'."""
            if len(w) < 8: return False
            for i in range(3, len(w)-3):
                if w[:i] in known_set and w[i:] in known_set:
                    return True
                # met tussenvoegsel -s- (arbeidSmarkt)
                if w[:i] in known_set and w[i:i+1] == "s" and w[i+1:] in known_set:
                    return True
            return False

        def smart_suggest(w, known_set, max_n=8):
            """Suggesties: edit-distance + fonetische score bovenaan."""
            base_sug = VaultManager._suggest(w, known_set, max_n=max_n*2)
            # Fonetische overeenkomst: begin + eind bewaren scoort hoger
            def phonetic_score(c):
                prefix = sum(1 for a,b in zip(w, c) if a==b)
                suffix = sum(1 for a,b in zip(reversed(w), reversed(c)) if a==b)
                lendiff = abs(len(c) - len(w))
                return (-prefix - suffix//2 + lendiff)
            return sorted(base_sug, key=phonetic_score)[:max_n]

        results = {}
        for raw in words:
            w = raw.strip("'-.,:!?;\"()[]{}").lower()
            if not w or len(w) < 2:
                results[raw] = {"spell": True}; continue
            if always_ok.search(raw):
                results[raw] = {"spell": True}; continue
            if any(c.isdigit() for c in w):
                results[raw] = {"spell": True}; continue

            # Hoofdletterwoorden: controleer de lowercase variant
            # (niet meer blind True — een typfout met hoofdletter is nog steeds fout)
            w_check = w  # altijd lowercase vergelijken

            if w_check in vault_ok or w_check in known:
                results[raw] = {"spell": True}
                # Indexeer ook op lowercase zodat frontend altijd kan vinden
                if raw != w_check:
                    results[w_check] = {"spell": True}
                continue

            # NL samengestelde woorden
            if lang == "nl" and is_compound_ok(w_check, known | vault_ok):
                results[raw] = {"spell": True}
                if raw != w_check: results[w_check] = {"spell": True}
                continue

            # Alleen-hoofdletter eerste letter EN >4 tekens → waarschijnlijk eigennaam
            if raw[0].isupper() and raw[1:].islower() and len(raw) > 4:
                results[raw] = {"spell": True}
                continue

            # Fout — genereer slimme suggesties
            sug = smart_suggest(w_check, known | vault_ok, max_n=8)
            results[raw] = {"spell": False, "suggestions": sug}
            # Dubbel indexeren: ook op lowercase zodat frontend altijd vindt
            if raw != w_check:
                results[w_check] = {"spell": False, "suggestions": sug}
        return results

    def grammar_check_lines(self, lines: list, lang: str = "en") -> list:
        """Grammaticacontrole per regel via patroonherkenning.
        Geeft [{row, col, len, msg, type}] met type 'grammar' of 'style'."""
        import re as _re
        errors = []
        if lang == "en":
            patterns = [
                (r"\b(\w+)\s+\1\b",
                 "dubbel woord", "grammar"),
                (r"\ba ([aeiouAEIOU]\w)",
                 "gebruik \'an\' voor klinkerwoorden (a \u2192 an)", "grammar"),
                (r"\ban ([bcdfghjklmnpqrstvwxyzBCDFGHJKLMNPQRSTVWXYZ]\w)",
                 "gebruik \'a\' voor medeklinkerwoorden (an \u2192 a)", "grammar"),
                (r"\b(more|less|greater|fewer|better|worse|higher|lower)\s+then\b",
                 "\'then\' is tijdsvolgorde; bedoel je \'than\' (vergelijking)?", "grammar"),
                (r"\byour\s+(are|were|going|coming|doing|getting|being)\b",
                 "\'your\' is bezittelijk; bedoel je \'you\'re\' (you are)?", "grammar"),
                (r"\bits\s+(a|an|the|very|quite|really|so|too)\b",
                 "\'its\' is bezittelijk; bedoel je \'it\'s\' (it is)?", "grammar"),
                (r"\btheir\s+(is|are|was|were|will|would)\b",
                 "bedoel je \'there\' of \'they\'re\'?", "grammar"),
                (r"\b(he|she|it)\s+(have|do|go|come|run|make|take|get)\b",
                 "derde persoon enkelvoud mist -s (has/does/goes/comes...)", "grammar"),
                (r"\bcould of\b|\bwould of\b|\bshould of\b",
                 "gebruik \'could/would/should have\', niet \'of\'", "grammar"),
            ]
        else:  # nl
            patterns = [
                (r"\b(\w+)\s+\1\b",
                 "dubbel woord", "grammar"),
                (r"\b(hij|zij|ze|het|jij|u)\s+(word|vind|houd|gebruik|verwacht|krijg|stop|begin)\b",
                 "dt-fout: derde persoon enkelvoud vereist -t (wordt/vindt/houdt\u2026)", "grammar"),
                (r"\bword\s+(door|niet|ook|wel|snel|vaak|al|nu|dan|hier|daar)\b",
                 "dt-fout: \'word\' \u2192 \'wordt\' (derde persoon enkelvoud)", "grammar"),
                (r"\bhun\s+(is|zijn|heeft|hebben|was|waren|loopt|werkt|gaat|doet|zegt)\b",
                 "\'hun\' als onderwerp is incorrect; gebruik \'zij\'", "grammar"),
                (r"\bhen\s+(is|zijn|heeft|hebben|was|waren)\b",
                 "\'hen\' als onderwerp is incorrect; gebruik \'zij\'", "grammar"),
                (r"\bhet\s+(man|vrouw|jongen|tafel|stoel|straat|stad|dag|nacht|school|kerk)\b",
                 "lidwoordfout: waarschijnlijk \'de\' in plaats van \'het\'", "grammar"),
                (r"\bde\s+(kind|boek|huis|woord|ding|getal|deel|recht|systeem|begin|einde|doel)\b",
                 "lidwoordfout: waarschijnlijk \'het\' in plaats van \'de\'", "grammar"),
                (r"\b(niet|nooit|niemand|nergens|niets)\b.{1,40}\b(niet|nooit|niemand|nergens|niets)\b",
                 "mogelijke dubbele ontkenning", "grammar"),
                (r"\bals\b.{1,40}\bals\b",
                 "tip: gebruik \'als\u2026dan\' bij voorwaardelijke zinnen", "style"),
            ]
        for row, line in enumerate(lines):
            for pat, msg, etype in patterns:
                for m in _re.finditer(pat, line, _re.IGNORECASE):
                    errors.append({
                        "row": row, "col": m.start(),
                        "len": m.end() - m.start(),
                        "msg": msg, "type": etype,
                    })
        return errors

    def _write_json(self, path, data):
        """Schrijf JSON atomisch via temp-bestand + rename.
        Voorkomt corruptie bij cloud-sync (OneDrive/Google Drive) die
        het bestand tijdens schrijven kan inlezen."""
        import tempfile as _tf, os as _os
        text = json.dumps(data, ensure_ascii=False, indent=2)
        # Schrijf naar temp-bestand in dezelfde map (zelfde filesystem = atomische rename)
        dir_ = path.parent
        try:
            fd, tmp = _tf.mkstemp(dir=str(dir_), suffix=".tmp")
            try:
                with _os.fdopen(fd, "w", encoding="utf-8") as f:
                    f.write(text)
                # Atomische rename — vervangt bestaand bestand in één operatie
                _os.replace(tmp, str(path))
            except Exception:
                try: _os.unlink(tmp)
                except: pass
                raise
        except OSError:
            # Fallback: directe schrijf (bijv. als temp-map op ander FS zit)
            path.write_text(text, encoding="utf-8")
    @property
    def path_str(self): return str(self.vault)


class ZKHandler(BaseHTTPRequestHandler):
    vault:   VaultManager = None
    verbose: bool         = False
    offline: bool         = False

    def log_message(self, fmt, *args):
        if self.verbose:
            print(f"  {self.command:<7} {args[1] if len(args)>1 else '???'}  {self.path}", flush=True)

    def _send(self, code, body, ct="application/json"):
        if isinstance(body,(dict,list)): body=json.dumps(body,ensure_ascii=False).encode("utf-8")
        elif isinstance(body,str):       body=body.encode("utf-8")
        try:
            self.send_response(code)
            self.send_header("Content-Type",ct)
            self.send_header("Content-Length",len(body))
            self.send_header("Access-Control-Allow-Origin","*")
            self.end_headers()
            self.wfile.write(body)
        except (BrokenPipeError, ConnectionResetError):
            pass  # client heeft verbinding verbroken — geen actie nodig

    def _sse_write(self, data: str) -> bool:
        """Schrijf een SSE-regel. Geeft False terug als de verbinding verbroken is."""
        try:
            self.wfile.write(data.encode("utf-8"))
            self.wfile.flush()
            return True
        except (BrokenPipeError, ConnectionResetError):
            return False

    def log_message(self, fmt, *args):
        pass  # Schakel standaard request-logging uit

    def handle_error(self, request, client_address):
        pass  # Onderdruk BrokenPipeError stacktraces in de console

    def _body(self):
        l=int(self.headers.get("Content-Length",0))
        return json.loads(self.rfile.read(l)) if l else {}

    def _raw_body(self):
        l=int(self.headers.get("Content-Length",0))
        return self.rfile.read(l) if l else b""

    def do_OPTIONS(self):
        self.send_response(200)
        for h,v in [("Access-Control-Allow-Origin","*"),
                    ("Access-Control-Allow-Methods","GET,POST,PUT,DELETE,OPTIONS"),
                    ("Access-Control-Allow-Headers","Content-Type")]:
            self.send_header(h,v)
        self.end_headers()

    def do_GET(self):
        p = urlparse(self.path).path.rstrip("/") or "/"
        if p=="/api/notes":            return self._send(200,self.vault.load_notes())
        if p=="/api/annotations":      return self._send(200,self.vault.load_annotations())
        if p=="/api/img-annotations":  return self._send(200,self.vault.load_img_annotations())
        if p=="/api/pdfs":             return self._send(200,self.vault.list_pdfs())
        if p=="/api/images":           return self._send(200,self.vault.list_images())
        if p=="/api/llm/models":  return self._llm_models()
        if p=="/api/config":      return self._send(200,{"vault_path":self.vault.path_str,"config":self.vault.get_config()})
        if p=="/api/api-keys":
            # Geeft terug welke providers geconfigureerd zijn + gemaskeerde waarde voor weergave
            result = {}
            for pr in ("anthropic","openai","google","openrouter","mistral"):
                key = self.vault.get_api_key(pr)
                result[pr] = {"set": bool(key), "preview": (key[:8]+"…"+key[-4:]) if len(key)>12 else ("●●●●" if key else "")}
            return self._send(200, result)
        if p=="/api/ext-pdfs":     return self._send(200,{"dirs":self.vault.get_ext_pdf_dirs(),"files":self.vault.scan_ext_pdfs()})
        if p.startswith("/api/browse"):
            from urllib.parse import parse_qs, urlparse as _up
            qs = parse_qs(_up(self.path).query)
            dir_path = qs.get("path",[""])[0].strip()
            return self._send(200, self._browse(dir_path))
        if p.startswith("/api/pdf/"):
            fname=unquote(p[9:])
            fp=self.vault.get_pdf_path(fname)
            if fp: return self._send(200,fp.read_bytes(),"application/pdf")
            return self._send(404,{"error":"PDF niet gevonden"})
        if p.startswith("/api/image/"):
            fname=unquote(p[11:])
            fp=self.vault.get_image_path(fname)
            if fp:
                ext=fp.suffix.lower()
                return self._send(200,fp.read_bytes(),IMAGE_MIME.get(ext,"image/jpeg"))
            return self._send(404,{"error":"Afbeelding niet gevonden"})
        # Statische bestanden — index.html via template (online/offline)
        if p == "/":
            return self._send(200, render_index(self.offline), "text/html")
        if p == "/app.js":
            fp = STATIC_DIR / "app.js"
            if fp.exists(): return self._send(200, fp.read_bytes(), "application/javascript")
        # Modules-bestanden (SOLID refactor — stap 1+)
        if p.startswith("/modules/"):
            rel = p[9:]  # strip /modules/
            fp = STATIC_DIR / "modules" / rel
            if fp.exists() and fp.suffix == ".js":
                return self._send(200, fp.read_bytes(), "application/javascript")
        # Vendor-bestanden (alleen beschikbaar in offline modus)
        if p.startswith("/vendor/"):
            rel = p[8:]  # strip /vendor/
            fp = VENDOR_DIR / rel
            if fp.exists():
                ext = fp.suffix.lower()
                mime = {"js":"application/javascript","css":"text/css",
                        "woff2":"font/woff2","woff":"font/woff"}.get(ext.lstrip("."), "application/octet-stream")
                return self._send(200, fp.read_bytes(), mime)
        return self._send(404,{"error":"Niet gevonden"})

    def do_POST(self):
        p=urlparse(self.path).path.rstrip("/")
        if p=="/api/search":
            body = self._body()
            q    = body.get("query","").strip()
            if not q: return self._send(200, {"results":[],"query":q})
            try:
                results = self.vault.fuzzy_search(q, max_results=80)
                return self._send(200, {"results": results, "query": q})
            except Exception as e:
                return self._send(500, {"error": str(e), "results": []})
        if p=="/api/llm/link-reasons":
            # LLM beoordeelt top-kandidaten en geeft reden per link
            body       = self._body()
            content_q  = body.get("content","").strip()
            candidates = body.get("candidates",[])   # [{id, title, tags}]
            model      = body.get("model","")
            if not content_q or not candidates or not model:
                return self._send(400, {"error": "content, candidates en model vereist"})
            try:
                reasons = self._llm_link_reasons(content_q, candidates, model)
                return self._send(200, {"reasons": reasons})
            except Exception as e:
                return self._send(500, {"error": str(e), "reasons": []})
        if p=="/api/suggest-links":
            body    = self._body()
            content = body.get("content","").strip()
            note_id = body.get("note_id","")      # huidig notitie-ID (om zichzelf te skippen)
            top_n   = int(body.get("top_n", 12))
            try:
                results = self.vault.suggest_links(content, note_id, top_n)
                return self._send(200, {"suggestions": results})
            except Exception as e:
                return self._send(500, {"error": str(e), "suggestions": []})
        if p=="/api/fulltext":
            body = self._body()
            q    = body.get("query","").strip()
            if not q: return self._send(200, {"results":[],"query":q})
            try:
                results = self.vault.fulltext_search(q, max_results=50)
                return self._send(200, {"results": results, "query": q})
            except Exception as e:
                return self._send(500, {"error": str(e), "results": []})
        if p=="/api/spellcheck":
            # Accepteert {words:[str], lines:[str], lang:"en"|"nl"|"auto"}
            # Geeft terug: {spell:{word:{spell,grammar}}, grammar:[{row,col,len,msg,type}]}
            body  = self._body()
            words = body.get("words", [])
            lines = body.get("lines", [])
            lang  = body.get("lang", "en")
            if lang == "auto":
                # Automatische taaldetectie op basis van stopwoorden
                import re as _re
                text = " ".join(lines).lower()
                nl_score = len(_re.findall(r"\b(de|het|een|en|van|in|is|dat|dit|met|ik|je|we|ze|ook|aan|maar|voor|niet)\b", text))
                en_score = len(_re.findall(r"\b(the|and|for|that|with|this|are|was|have|from|they|will|been|not|can|but)\b", text))
                lang = "nl" if nl_score > en_score else "en"
            spell_results  = self.vault.spellcheck_words(words, lang)
            grammar_errors = self.vault.grammar_check_lines(lines, lang) if lines else []
            return self._send(200, {
                "spell":   spell_results,
                "grammar": grammar_errors,
                "lang":    lang,
            })
        if p=="/api/notes":
            return self._send(200,self.vault.save_note(self._body()))
        if p=="/api/annotations":
            a=self._body()
            if isinstance(a,list): self.vault.save_annotations(a); return self._send(200,{"ok":True})
            return self._send(400,{"error":"Verwacht een lijst"})
        if p=="/api/img-annotations":
            a=self._body()
            if isinstance(a,list): self.vault.save_img_annotations(a); return self._send(200,{"ok":True})
            return self._send(400,{"error":"Verwacht een lijst"})
        if p=="/api/pdfs":
            ct=self.headers.get("Content-Type","")
            if "multipart" in ct:
                boundary=ct.split("boundary=")[-1].encode()
                for part in self._raw_body().split(b"--"+boundary):
                    if b"filename=" in part:
                        hdr,_,body=part.partition(b"\r\n\r\n")
                        fname=hdr.split(b"filename=")[1].split(b'"')[1].decode()
                        saved=self.vault.save_pdf(fname,body.rstrip(b"\r\n--"))
                        return self._send(200,{"name":saved})
            return self._send(400,{"error":"Multipart vereist"})
        if p=="/api/images":
            ct=self.headers.get("Content-Type","")
            if "multipart" in ct:
                boundary=ct.split("boundary=")[-1].encode()
                for part in self._raw_body().split(b"--"+boundary):
                    if b"filename=" in part:
                        hdr,_,body=part.partition(b"\r\n\r\n")
                        fname=hdr.split(b"filename=")[1].split(b'"')[1].decode()
                        return self._send(200,self.vault.save_image(fname,body.rstrip(b"\r\n--")))
            return self._send(400,{"error":"Multipart vereist"})
        if p=="/api/llm/improve-text":   return self._llm_improve_text()
        if p=="/api/llm/suggest-tags":   return self._llm_suggest_tags()
        if p=="/api/llm/chat":           return self._llm_chat()
        if p=="/api/llm/similar":        return self._llm_similar()
        if p=="/api/llm/graphrag":       return self._llm_graphrag()
        if p=="/api/llm/summarize-pdf":  return self._llm_summarize_pdf()
        if p=="/api/llm/describe-image": return self._llm_describe_image()
        if p=="/api/llm/summarize-note": return self._llm_summarize_note()
        if p=="/api/llm/mindmap":        return self._llm_mindmap()
        if p=="/api/import-url":         return self._import_url()
        if p=="/api/import-docx":        return self._import_docx()
        if p=="/api/ext-pdfs":
            dirs = self._body().get("dirs", [])
            self.vault.set_ext_pdf_dirs(dirs)
            return self._send(200, {"ok": True, "dirs": dirs})
        if p=="/api/api-keys":
            body = self._body()
            for provider in ("anthropic","openai","google","openrouter","mistral"):
                if provider in body:
                    self.vault.set_api_key(provider, body[provider])
            status = {pr: bool(self.vault.get_api_key(pr)) for pr in ("anthropic","openai","google","openrouter","mistral")}
            return self._send(200, {"ok": True, "configured": status})
        if p=="/api/vault":
            body=self._body()
            np=body.get("path","").strip()
            if not np: return self._send(400,{"error":"Pad vereist"})
            ZKHandler.vault=VaultManager(Path(np))
            return self._send(200,{"vault_path":ZKHandler.vault.path_str})
        if p=="/api/config":
            body=self._body()
            allowed={"pdf_personal_use","pdf_personal_email","review_data"}
            update={k:v for k,v in body.items() if k in allowed}
            self.vault.save_config(update)
            return self._send(200,{"ok":True})
        if p=="/api/cleanup-vault":
            # Verwijder CSS-rommel uit alle opgeslagen notities (batch-fix)
            try:
                import re as _re
                notes_dir = self.vault.notes_dir
                css_pat = _re.compile(
                    r'#?[0-9a-fA-F]{0,8};?(?:[\w-]+:[^;">\
]{1,60};?){1,10}"?>'
                )
                cleaned = 0; skipped = 0; errors = []
                for md_file in notes_dir.glob("*.md"):
                    try:
                        raw = md_file.read_text(encoding="utf-8")
                        if raw.startswith("---"):
                            parts = raw.split("---", 2)
                            if len(parts) >= 3:
                                new_body = css_pat.sub("", parts[2])
                                import re as _re2
                                new_body = _re2.sub(r'^#+\s*$', '', new_body, flags=_re2.MULTILINE)
                                new_body = _re2.sub(r'\n{3,}', '\n\n', new_body)
                                if new_body != parts[2]:
                                    md_file.write_text("---" + parts[1] + "---" + new_body, encoding="utf-8")
                                    cleaned += 1
                                else:
                                    skipped += 1
                            else:
                                skipped += 1
                        else:
                            new_raw = css_pat.sub("", raw)
                            if new_raw != raw:
                                md_file.write_text(new_raw, encoding="utf-8")
                                cleaned += 1
                            else:
                                skipped += 1
                    except Exception as e:
                        errors.append({"file": md_file.name, "error": str(e)})
                self.vault._cache = None
                return self._send(200, {"ok": True, "cleaned": cleaned, "skipped": skipped, "errors": errors})
            except Exception as e:
                return self._send(500, {"ok": False, "error": str(e)})
        return self._send(404,{"error":"Niet gevonden"})

    def do_PUT(self):
        p=urlparse(self.path).path.rstrip("/")
        if p.startswith("/api/notes/"):
            note=self._body(); note["id"]=unquote(p[11:])
            return self._send(200,self.vault.save_note(note))
        return self._send(404,{"error":"Niet gevonden"})

    def do_DELETE(self):
        p=urlparse(self.path).path.rstrip("/")
        if p.startswith("/api/notes/"):
            return self._send(200,{"deleted":self.vault.delete_note(unquote(p[11:]))})
        if p.startswith("/api/images/"):
            return self._send(200,{"deleted":self.vault.delete_image(unquote(p[12:]))})
        if p.startswith("/api/pdfs/"):
            fname=unquote(p[10:])
            result=self.vault.delete_pdf(fname)
            return self._send(200,{"ok":True,"filename":fname,**result})
        return self._send(404,{"error":"Niet gevonden"})

    # ── LLM ────────────────────────────────────────────────────────────────────
    def _ollama(self): return os.environ.get("OLLAMA_URL","http://localhost:11434")

    def _ollama_post(self, endpoint, payload, timeout=180):
        req=urllib.request.Request(self._ollama()+endpoint,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type":"application/json"},method="POST")
        with urllib.request.urlopen(req,timeout=timeout) as r:
            return json.loads(r.read())

    def _llm_models(self):
        try:
            r=urllib.request.urlopen(self._ollama()+"/api/tags",timeout=3)
            d=json.loads(r.read())
            return self._send(200,{"models":[m["name"] for m in d.get("models",[])],"ok":True,"ollama_url":self._ollama()})
        except Exception as e:
            return self._send(200,{"models":[],"ok":False,"error":str(e),"ollama_url":self._ollama()})

    def _browse(self, path):
        """Bestandsverkenner: geeft inhoud van een map terug."""
        import os, platform
        if not path:
            # Roots: home + schijven afhankelijk van OS
            roots = []
            home = Path.home()
            roots.append({"name": "🏠 "+home.name, "path": str(home), "type": "dir"})
            if platform.system() == "Darwin":  # macOS
                volumes = Path("/Volumes")
                if volumes.exists():
                    for v in sorted(volumes.iterdir()):
                        if not v.name.startswith("."):
                            roots.append({"name": "💿 "+v.name, "path": str(v), "type": "dir"})
            elif platform.system() == "Linux":
                for r in ["/", "/mnt", "/media", "/home"]:
                    rp = Path(r)
                    if rp.exists() and str(rp) != str(home.parent):
                        roots.append({"name": "🖴 "+r, "path": r, "type": "dir"})
            elif platform.system() == "Windows":
                import string
                for letter in string.ascii_uppercase:
                    drive = Path(letter+":\\")
                    if drive.exists():
                        roots.append({"name": "💾 "+letter+":\\", "path": str(drive), "type": "dir"})
            return {"path": "", "parent": "", "items": roots, "is_root": True}

        p = Path(path)
        if not p.exists() or not p.is_dir():
            return {"error": "Pad niet gevonden: "+str(path), "path": path, "items": []}
        items = []
        try:
            import os, signal

            dirs, pdfs = [], []
            # Harde limiet: max 2 seconden, max 500 entries
            deadline = __import__("time").time() + 2.0
            with os.scandir(str(p)) as it:
                for e in it:
                    if __import__("time").time() > deadline:
                        break                        # stop bij timeout
                    if e.name.startswith("."): continue
                    try:
                        if e.is_dir(follow_symlinks=False):
                            dirs.append({"name": e.name, "path": e.path, "type": "dir"})
                        elif e.name.lower().endswith(".pdf"):
                            try: sz = e.stat(follow_symlinks=False).st_size
                            except: sz = 0
                            pdfs.append({"name": e.name, "path": e.path,
                                         "type": "pdf", "size": sz})
                    except (PermissionError, OSError):
                        pass
                    if len(dirs) + len(pdfs) >= 500:
                        break

            dirs.sort(key=lambda x: x["name"].lower())
            pdfs.sort(key=lambda x: x["name"].lower())
            items = dirs + pdfs

        except (PermissionError, OSError) as ex:
            return {"error": "Geen toegang: "+str(ex), "path": path, "items": []}
        parent = str(p.parent) if p.parent != p else ""
        return {"path": str(p), "parent": parent, "items": items}

    def _llm_improve_text(self):
        """Stuur tekst naar LLM voor taalverbetering."""
        body  = self._body()
        text  = body.get("text", "").strip()
        lang  = body.get("lang", "nl")
        model = body.get("model", "")
        if not text or not model:
            return self._send(400, {"error": "text en model zijn vereist"})
        lang_label = "Nederlands" if lang == "nl" else "English"
        prompt = (
            f"Je bent een taalkundige redacteur. Verbeter de volgende tekst in het {lang_label}. "
            f"Geef ALLEEN de verbeterde tekst terug, zonder uitleg, zonder commentaar, "
            f"zonder markdown-opmaak rondom de tekst zelf. "
            f"Behoud de oorspronkelijke structuur en alinea-indeling.\n\n"
            f"Originele tekst:\n{text}"
        )
        try:
            import urllib.request as _req, json as _json
            payload = _json.dumps({
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "options": {"temperature": 0.3, "num_predict": 800}
            }).encode()
            req = _req.Request(
                "http://localhost:11434/api/chat",
                data=payload, headers={"Content-Type": "application/json"}, method="POST"
            )
            with _req.urlopen(req, timeout=60) as r:
                data = _json.loads(r.read())
            improved = (data.get("message", {}).get("content", "") or "").strip()
            if not improved:
                return self._send(500, {"error": "Geen respons van LLM"})
            return self._send(200, {"improved": improved, "original": text})
        except Exception as e:
            return self._send(500, {"error": str(e)})

    def _llm_suggest_tags(self):
        """Stel slimme tags voor bij een notitie — gebruikt bestaande vault-tags semantisch.
        Geeft gewone JSON terug (geen SSE)."""
        body        = self._body()
        content     = body.get("content", "").strip()
        model       = body.get("model", "")
        current     = body.get("current_tags", [])
        all_tags    = body.get("all_tags", [])

        if not content or not model:
            return self._send(400, {"error": "content en model zijn vereist"})

        # Bouw een frequentie-overzicht van bestaande tags (meest gebruikt = meest relevant)
        # Haal ook notitie-titels op voor contextuele overlap
        try:
            notes = self.vault.load_notes()
            tag_freq = {}
            for n in notes:
                for t in (n.get("tags") or []):
                    tag_freq[t] = tag_freq.get(t, 0) + 1
            # Sorteer op frequentie — meest gebruikte tags bovenaan
            sorted_tags = sorted(tag_freq.items(), key=lambda x: -x[1])
            top_tags = [t for t, _ in sorted_tags[:80]]
        except Exception:
            top_tags = all_tags[:60]

        # Bepaal welke bestaande tags ook echt in de tekst voorkomen (directe match)
        text_lower = content.lower()
        matching_tags = [t for t in top_tags
                         if t.replace("_"," ") in text_lower or t in text_lower][:20]

        existing_str = ", ".join(top_tags[:60]) if top_tags else ""
        matching_str = ", ".join(matching_tags) if matching_tags else "geen directe overeenkomsten"

        prompt = (
            "Je bent een Zettelkasten-assistent. Stel maximaal 6 tags voor bij de onderstaande tekst.\n\n"
            "REGELS:\n"
            "- Alleen lowercase, enkelvoud, geen spaties (gebruik underscore), geen #\n"
            "- HERGEBRUIK bestaande tags waar mogelijk — dit verbindt notities\n"
            "- Voeg maximaal 2 nieuwe tags toe als bestaande tags niet passend zijn\n"
            "- Kies tags op conceptueel niveau (niet te specifiek, niet te breed)\n\n"
            f"BESTAANDE TAGS IN DE VAULT (meest gebruikt eerst):\n{existing_str}\n\n"
            f"TAGS DIE DIRECT IN DE TEKST VOORKOMEN (sterke kandidaten):\n{matching_str}\n\n"
            f"HUIDIGE TAGS VAN DEZE NOTITIE (niet herhalen):\n{', '.join(current) if current else 'geen'}\n\n"
            "Antwoord ALLEEN met een JSON-array, bijv: [\"concept\",\"ai\",\"methode\"]\n\n"
            f"TEKST:\n{content[:4000]}"
        )

        try:
            text = ""

            if model.startswith("claude"):
                api_key = self.vault.get_api_key("anthropic")
                if not api_key:
                    return self._send(401, {"error": "Anthropic API-sleutel niet ingesteld"})
                payload = json.dumps({
                    "model": model, "max_tokens": 200,
                    "messages": [{"role": "user", "content": prompt}]
                }).encode()
                req = urllib.request.Request(
                    "https://api.anthropic.com/v1/messages", data=payload,
                    headers={"Content-Type": "application/json",
                             "x-api-key": api_key,
                             "anthropic-version": "2023-06-01"}, method="POST")
                with urllib.request.urlopen(req, timeout=30) as r:
                    d = json.loads(r.read())
                text = d.get("content", [{}])[0].get("text", "")

            elif model.startswith("gemini"):
                api_key = self.vault.get_api_key("google")
                if not api_key:
                    return self._send(401, {"error": "Google API-sleutel niet ingesteld"})
                payload = json.dumps({
                    "contents": [{"role": "user", "parts": [{"text": prompt}]}],
                    "generationConfig": {"maxOutputTokens": 200, "temperature": 0.2}
                }).encode()
                url2 = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
                req = urllib.request.Request(url2, data=payload,
                    headers={"Content-Type": "application/json"}, method="POST")
                with urllib.request.urlopen(req, timeout=30) as r:
                    d = json.loads(r.read())
                text = d.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "")

            elif model.startswith("gpt") or model.startswith("o1") or model.startswith("o3") or model.startswith("o4"):
                api_key = self.vault.get_api_key("openai")
                if not api_key:
                    return self._send(401, {"error": "OpenAI API-sleutel niet ingesteld"})
                payload = json.dumps({
                    "model": model, "max_tokens": 200,
                    "messages": [{"role": "user", "content": prompt}]
                }).encode()
                req = urllib.request.Request(
                    "https://api.openai.com/v1/chat/completions", data=payload,
                    headers={"Content-Type": "application/json",
                             "Authorization": "Bearer " + api_key}, method="POST")
                with urllib.request.urlopen(req, timeout=30) as r:
                    d = json.loads(r.read())
                text = d.get("choices", [{}])[0].get("message", {}).get("content", "")

            elif model.startswith("mistral") or "/" in model:  # Mistral direct of OpenRouter
                key_provider2 = "mistral" if model.startswith("mistral") else "openrouter"
                base_url2 = "https://api.mistral.ai/v1/chat/completions" if key_provider2=="mistral" else "https://openrouter.ai/api/v1/chat/completions"
                api_key = self.vault.get_api_key(key_provider2)
                if not api_key:
                    return self._send(401, {"error": "OpenRouter API-sleutel niet ingesteld"})
                payload = json.dumps({
                    "model": model, "max_tokens": 200,
                    "messages": [{"role": "user", "content": prompt}]
                }).encode()
                req = urllib.request.Request(
                    "https://openrouter.ai/api/v1/chat/completions", data=payload,
                    headers={"Content-Type": "application/json",
                             "Authorization": "Bearer " + api_key,
                             "HTTP-Referer": "http://localhost:8899",
                             "X-Title": "Zettelkasten"}, method="POST")
                with urllib.request.urlopen(req, timeout=30) as r:
                    d = json.loads(r.read())
                text = d.get("choices", [{}])[0].get("message", {}).get("content", "")

            elif model.startswith("mistral") or model.startswith("magistral") or model.startswith("ministral"):
                api_key = self.vault.get_api_key("mistral")
                if not api_key:
                    return self._send(401, {"error": "Mistral API-sleutel niet ingesteld"})
                payload = json.dumps({
                    "model": model, "max_tokens": 200,
                    "messages": [{"role": "user", "content": prompt}]
                }).encode()
                req = urllib.request.Request(
                    "https://api.mistral.ai/v1/chat/completions", data=payload,
                    headers={"Content-Type": "application/json",
                             "Authorization": "Bearer " + api_key}, method="POST")
                with urllib.request.urlopen(req, timeout=30) as r:
                    d = json.loads(r.read())
                text = d.get("choices", [{}])[0].get("message", {}).get("content", "")

            else:
                # Lokaal Ollama
                payload = json.dumps({
                    "model": self._best_local_model(model),
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                    "options": {"temperature": 0.2, "num_predict": 200}
                }).encode()
                req = urllib.request.Request(
                    self._ollama() + "/api/chat", data=payload,
                    headers={"Content-Type": "application/json"}, method="POST")
                with urllib.request.urlopen(req, timeout=60) as r:
                    d = json.loads(r.read())
                text = d.get("message", {}).get("content", "")

            # Extraheer JSON-array uit de response
            import re as _re2
            m = _re2.search(r'\[.*?\]', text, _re2.DOTALL)
            if not m:
                return self._send(200, {"tags": []})
            tags = json.loads(m.group(0))
            # Valideer en normaliseer
            tags = [str(t).strip().lower().replace(" ","_").lstrip("#")
                    for t in tags if isinstance(t,str) and t.strip()]
            tags = [t for t in tags if t and t not in current][:8]
            return self._send(200, {"tags": tags})

        except Exception as e:
            return self._send(500, {"error": str(e)})
    def _build_context(self, body):
        """Bouw kenniscontext op vanuit notities, PDFs en afbeeldingen.
        Bewaakt een totaal-budget van ~60.000 tekens om 400-fouten te voorkomen."""
        ctx_notes  = body.get("context_notes", [])
        ctx_pdfs   = body.get("context_pdfs", [])
        ctx_imgs   = body.get("context_images", [])
        ctx_ext    = body.get("context_ext_pdfs", [])

        BUDGET = 60_000   # max tekens over alle context-blokken
        used   = 0
        parts  = []

        def add(text):
            nonlocal used
            remaining = BUDGET - used
            if remaining <= 0:
                return False
            chunk = text[:remaining]
            parts.append(chunk)
            used += len(chunk)
            return used < BUDGET

        if ctx_notes:
            notes = [x for x in self.vault.load_notes() if x["id"] in ctx_notes]
            # Sorteer op relevantie: notities met meer content eerst
            notes.sort(key=lambda n: len(n.get("content","") or ""), reverse=True)
            for n in notes[:20]:   # max 20 notities, budget bepaalt hoeveel er echt in passen
                snippet = (n.get("content") or "")[:3000]
                if not add(f"## Notitie: {n['title']}\n{snippet}"):
                    break

        if ctx_pdfs and used < BUDGET:
            annots = self.vault.load_annotations()
            for pn in ctx_pdfs[:4]:
                pa = [a for a in annots if a.get("file") == pn]
                if pa:
                    lines = ['- "'+a["text"]+'"'+(("\n  Noot: "+a["note"]) if a.get("note") else "") for a in pa[:20]]
                    if not add(f"## PDF annotaties: {pn}\n"+"\n".join(lines)):
                        break
                txt = self.vault.extract_pdf_text(pn, 4000) or ""
                if txt.strip():
                    if not add(f"## PDF tekst: {pn}\n{txt}"):
                        break

        if ctx_ext and used < BUDGET:
            for ep in ctx_ext[:4]:
                txt  = self.vault.extract_pdf_text_from_path(ep, 4000) or ""
                name = Path(ep).name
                if not add(f"## Extern PDF: {name}\n"+(txt.strip() or "(geen tekst)")):
                    break

        if ctx_imgs:
            img_lines = [f"- {n}" for n in ctx_imgs[:6]]
            if img_lines:
                add("## Afbeeldingen:\n"+"\n".join(img_lines))

        return parts

    def _llm_chat(self):
        body=self._body()
        raw_model=body.get("model","")
        messages=body.get("messages",[])

        parts=self._build_context(body)
        system=("Je bent een behulpzame kennisassistent voor een Zettelkasten. "
                "Antwoord in de taal van de gebruiker. Wees analytisch en precies.")
        if parts: system+="\n\n# Kenniscontext:\n\n"+"\n\n---\n\n".join(parts)

        self.send_response(200)
        self.send_header("Content-Type","text/event-stream")
        self.send_header("Cache-Control","no-cache")
        self.send_header("Access-Control-Allow-Origin","*")
        self.end_headers()

        if raw_model.startswith("claude"):
            self._stream_anthropic(raw_model, system, messages)
        elif raw_model.startswith("gemini"):
            self._stream_google(raw_model, system, messages)
        elif raw_model.startswith("gpt") or raw_model.startswith("o1") or raw_model.startswith("o3") or raw_model.startswith("o4"):
            self._stream_openai(raw_model, system, messages)
        elif "/" in raw_model:  # OpenRouter formaat: "org/model" (incl. Kimi)
            self._stream_openrouter(raw_model, system, messages)
        elif (raw_model.startswith("mistral") or raw_model.startswith("magistral")
              or raw_model.startswith("ministral")):
            self._stream_mistral(raw_model, system, messages)
        else:
            self._stream_ollama(self._best_local_model(raw_model), system, messages)

    # ── Semantische similariteit (TF-IDF cosine, geen externe deps) ─────────────
    def _tfidf_vectors(self, docs):
        """Berekent TF-IDF vectoren voor een lijst teksten. Geeft (vocab, matrix) terug."""
        import math, re
        stopwords = {"de","het","een","van","in","is","op","en","te","dat","die","zijn","met",
                     "voor","aan","er","maar","om","dan","als","ook","dit","nog","wel","ze",
                     "the","a","an","of","to","in","is","it","and","for","on","with","that"}
        def tokenize(t):
            return [w for w in re.findall(r"[a-z\u00c0-\u024f]{3,}", t.lower()) if w not in stopwords]
        tokenized = [tokenize(d) for d in docs]
        vocab = {w for toks in tokenized for w in toks}
        vocab = sorted(vocab)
        vidx  = {w:i for i,w in enumerate(vocab)}
        N     = len(docs)
        # IDF
        df = {}
        for toks in tokenized:
            for w in set(toks):
                df[w] = df.get(w,0)+1
        idf = {w: math.log((N+1)/(df.get(w,0)+1))+1 for w in vocab}
        # TF-IDF matrix (sparse as list of dicts)
        vecs = []
        for toks in tokenized:
            tf = {}
            for w in toks: tf[w] = tf.get(w,0)+1
            n = len(toks) or 1
            vec = {vidx[w]: (tf[w]/n)*idf[w] for w in tf if w in vidx}
            norm = math.sqrt(sum(v*v for v in vec.values())) or 1
            vecs.append({k: v/norm for k,v in vec.items()})
        return vecs

    def _cosine(self, a, b):
        return sum(a.get(k,0)*v for k,v in b.items())

    def _llm_similar(self):
        """Geeft top-N semantisch vergelijkbare notities terug via TF-IDF cosine."""
        body   = self._body()
        note_id= body.get("note_id","")
        top_n  = int(body.get("top_n", 6))
        notes  = self.vault.load_notes()
        if len(notes) < 2:
            return self._send(200, {"similar": []})
        texts  = [n.get("title","")+" "+n.get("content","") for n in notes]
        vecs   = self._tfidf_vectors(texts)
        idx    = next((i for i,n in enumerate(notes) if n["id"]==note_id), None)
        if idx is None:
            return self._send(200, {"similar": []})
        qvec   = vecs[idx]
        scores = [(self._cosine(qvec, vecs[i]), notes[i]) for i in range(len(notes)) if i!=idx]
        scores.sort(key=lambda x: -x[0])
        # Filter op minimale score (> 0.05) en geef titels + scores terug
        result = [{"id":n["id"],"title":n["title"],"score":round(s,3)}
                  for s,n in scores[:top_n] if s > 0.05]
        return self._send(200, {"similar": result})

    # ── GraphRAG: graaf-bewuste vragen over de kennisbasis ───────────────────────
    def _llm_graphrag(self):
        """GraphRAG: beantwoordt een vraag met graaf-context (notitie + buren + communities).
        Gebruikt het geselecteerde model (Anthropic/Google/OpenAI/OpenRouter/Ollama)."""
        import re as _re, math
        body     = self._body()
        question = body.get("question","")
        model    = body.get("model","")
        top_n    = int(body.get("top_n", 5))
        notes    = self.vault.load_notes()
        if not notes or not question:
            return self._send(400, {"error":"Geen vraag of notities"})

        # 1. Vind top-N relevante notities via TF-IDF
        texts  = [n.get("title","")+" "+n.get("content","") for n in notes]
        vecs   = self._tfidf_vectors([question] + texts)
        qvec   = vecs[0]
        scores = [(self._cosine(qvec, vecs[i+1]), notes[i]) for i in range(len(notes))]
        scores.sort(key=lambda x: -x[0])
        top_notes = [n for _,n in scores[:top_n] if _ > 0.02]
        if not top_notes:
            top_notes = [n for _,n in scores[:3]]  # altijd minimaal 3

        # 2. Bouw graaf-context: voor elke top-notitie, voeg ook directe buren toe
        import re as re2
        def extract_links(content):
            return [m.group(1) for m in re2.finditer(r'\[\[([^\]]+)\]\]', content or "")]

        note_map = {n["id"]: n for n in notes}
        context_ids = set(n["id"] for n in top_notes)
        for n in top_notes:
            for lid in extract_links(n.get("content","")):
                if lid in note_map: context_ids.add(lid)
            # ook backlinks
            for other in notes:
                if n["id"] in extract_links(other.get("content","")):
                    context_ids.add(other["id"])

        context_notes = [note_map[nid] for nid in context_ids if nid in note_map]

        # 3. Bouw community-samenvatting via label propagation (lightweight)
        def community_label(note_ids, note_map):
            labels = {nid: nid for nid in note_ids}
            for _ in range(4):
                for nid in note_ids:
                    n = note_map.get(nid)
                    if not n: continue
                    neighbors = [l for l in extract_links(n.get("content","")) if l in note_ids]
                    backs = [other for other in note_ids
                             if nid in extract_links((note_map.get(other) or {}).get("content",""))]
                    all_nb = neighbors + backs
                    if not all_nb: continue
                    votes = {}
                    for nb in all_nb:
                        lbl = labels[nb]
                        votes[lbl] = votes.get(lbl,0)+1
                    labels[nid] = max(votes, key=votes.get)
            return labels

        all_ids = [n["id"] for n in notes]
        labels  = community_label(all_ids, note_map)
        # Groepeer context_notes per community
        comm_groups = {}
        for nid in context_ids:
            lbl = labels.get(nid, nid)
            comm_groups.setdefault(lbl, []).append(nid)

        # 4. Stel systeem-prompt samen met graaf-context
        parts = []
        parts.append(f"## Vraag van de gebruiker\n{question}")
        parts.append(f"## Graaf-context ({len(context_notes)} notities, {len(comm_groups)} communities)")

        for lbl, members in sorted(comm_groups.items(), key=lambda x:-len(x[1]))[:6]:
            hub = sorted(members, key=lambda nid: sum(
                1 for x in notes if nid in extract_links(x.get("content",""))
            ), reverse=True)[0]
            hub_note = note_map.get(hub)
            member_titles = [note_map[nid]["title"] for nid in members if nid in note_map]
            parts.append(f"### Community (hub: {hub_note['title'] if hub_note else hub})\n"
                        f"Leden: {', '.join(member_titles[:8])}")

        parts.append("## Relevante notities (volledig)")
        for n in context_notes[:8]:
            links_out = extract_links(n.get("content",""))
            link_titles = [note_map[l]["title"] for l in links_out if l in note_map]
            backlinks   = [x["title"] for x in notes if n["id"] in extract_links(x.get("content",""))]
            parts.append(
                f"### {n['title']}\n"
                f"Tags: {', '.join(n.get('tags',[]) or ['–'])}\n"
                f"Links naar: {', '.join(link_titles) or '–'}\n"
                f"Backlinks: {', '.join(backlinks[:5]) or '–'}\n\n"
                f"{(n.get('content') or '')[:2000]}"
            )

        system_intro = (
            "Je bent een geavanceerde kennisassistent die werkt met een Zettelkasten knowledge graph. "
            "Je hebt toegang tot de structuur van het kennisnetwerk: welke notities verwant zijn, "
            "welke communities er bestaan, en hoe ideeën met elkaar verbonden zijn. "
            "Gebruik deze graaf-structuur actief in je antwoord: verwijs naar verbindingen, "
            "wijs op clusters van ideeën, en signaleer eventuele kennishiaten. "
            "Antwoord in de taal van de gebruiker. Wees analytisch en concreet."
        )
        # Budgetbewaking: max 50.000 tekens voor de context-delen
        BUDGET = 50_000
        context_text = "\n\n---\n\n".join(parts)
        if len(context_text) > BUDGET:
            context_text = context_text[:BUDGET] + "\n\n[...context ingekort om limiet te respecteren]"
        system = system_intro + "\n\n" + context_text
        messages = [{"role":"user","content":question}]

        self.send_response(200)
        self.send_header("Content-Type","text/event-stream")
        self.send_header("Cache-Control","no-cache")
        self.send_header("Access-Control-Allow-Origin","*")
        self.end_headers()

        if model.startswith("claude"):
            self._stream_anthropic(model, system, messages)
        elif model.startswith("gemini"):
            self._stream_google(model, system, messages)
        elif model.startswith("gpt") or model.startswith("o1") or model.startswith("o3") or model.startswith("o4"):
            self._stream_openai(model, system, messages)
        elif "/" in model:
            self._stream_openrouter(model, system, messages)
        elif (model.startswith("mistral") or model.startswith("magistral")
              or model.startswith("ministral")):
            self._stream_mistral(model, system, messages)
        else:
            self._stream_ollama(self._best_local_model(model), system, messages)

    def _stream_anthropic(self, model, system, messages):
        """Streaming via Anthropic API. Sleutel via instellingen of ANTHROPIC_API_KEY env."""
        api_key = self.vault.get_api_key("anthropic")
        if not api_key:
            try: self.wfile.write(("data: "+json.dumps({"error":"Anthropic API-sleutel niet ingesteld — voeg toe via ⚙ Instellingen → API-sleutels, of: export ANTHROPIC_API_KEY=sk-ant-..."})+"\n\n").encode()); self.wfile.flush()
            except: pass
            return
        # Alleen user/assistant roles, geen lege content, alternerende volgorde afdwingen
        ant_msgs = []
        for m in messages:
            role = m.get("role","")
            content = (m.get("content") or "").strip()
            if role not in ("user","assistant") or not content:
                continue
            # Anthropic vereist alternerende user/assistant; voorkom dubbele rollen
            if ant_msgs and ant_msgs[-1]["role"] == role:
                ant_msgs[-1]["content"] += "\n" + content
            else:
                ant_msgs.append({"role": role, "content": content})
        # Moet eindigen op user; voeg placeholder toe als leeg of eindigt op assistant
        if not ant_msgs or ant_msgs[-1]["role"] != "user":
            ant_msgs.append({"role": "user", "content": "(Vervolg)"})

        payload={"model":model,"max_tokens":4096,"system":system,"messages":ant_msgs,"stream":True}
        try:
            req=urllib.request.Request("https://api.anthropic.com/v1/messages",
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type":"application/json","x-api-key":api_key,"anthropic-version":"2023-06-01"},
                method="POST")
            with urllib.request.urlopen(req,timeout=120) as resp:
                for line in resp:
                    line=line.strip()
                    if not line: continue
                    try:
                        ls=line.decode("utf-8") if isinstance(line,bytes) else line
                        if ls.startswith("data: "):
                            ev=json.loads(ls[6:])
                            t=ev.get("type","")
                            if t=="content_block_delta":
                                delta=ev.get("delta",{}).get("text","")
                                self.wfile.write(("data: "+json.dumps({"delta":delta,"done":False})+"\n\n").encode("utf-8"))
                                self.wfile.flush()
                            elif t=="message_stop":
                                self.wfile.write(("data: "+json.dumps({"delta":"","done":True})+"\n\n").encode("utf-8"))
                                self.wfile.flush(); break
                    except: pass
        except urllib.error.HTTPError as e:
            try:
                body = e.read().decode("utf-8","replace")
                detail = json.loads(body).get("error",{}).get("message", body[:300])
            except: detail = str(e)
            try: self.wfile.write(("data: "+json.dumps({"error":f"Anthropic {e.code}: {detail}"})+"\n\n").encode()); self.wfile.flush()
            except: pass
        except Exception as e:
            try: self.wfile.write(("data: "+json.dumps({"error":"Anthropic API: "+str(e)})+"\n\n").encode()); self.wfile.flush()
            except: pass

    def _stream_google(self, model, system, messages):
        """Streaming via Google Gemini API. Sleutel via instellingen of GOOGLE_API_KEY env."""
        api_key = self.vault.get_api_key("google")
        if not api_key:
            try: self.wfile.write(("data: "+json.dumps({"error":"Google API-sleutel niet ingesteld — voeg toe via ⚙ Instellingen → API-sleutels, of: export GOOGLE_API_KEY=..."})+"\n\n").encode()); self.wfile.flush()
            except: pass
            return
        # Bouw Gemini messages op (system als eerste user-turn)
        gem_msgs = [{"role":"user","parts":[{"text":"Systeeminstructie: "+system}]},
                    {"role":"model","parts":[{"text":"Begrepen."}]}]
        for m in messages:
            role = "user" if m["role"]=="user" else "model"
            gem_msgs.append({"role":role,"parts":[{"text":m["content"]}]})
        payload = {"contents": gem_msgs,
                   "generationConfig":{"maxOutputTokens":4096}}
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:streamGenerateContent?alt=sse&key={api_key}"
        try:
            req = urllib.request.Request(url,
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type":"application/json"}, method="POST")
            with urllib.request.urlopen(req, timeout=120) as resp:
                for line in resp:
                    line = line.strip()
                    if not line: continue
                    try:
                        ls = line.decode("utf-8") if isinstance(line,bytes) else line
                        if ls.startswith("data: "):
                            ev = json.loads(ls[6:])
                            delta = ""
                            for cand in ev.get("candidates",[]):
                                for part in cand.get("content",{}).get("parts",[]):
                                    delta += part.get("text","")
                            done = ev.get("candidates",[{}])[-1].get("finishReason","") != ""
                            self.wfile.write(("data: "+json.dumps({"delta":delta,"done":done})+"\n\n").encode("utf-8"))
                            self.wfile.flush()
                    except: pass
                self.wfile.write(("data: "+json.dumps({"delta":"","done":True})+"\n\n").encode("utf-8"))
                self.wfile.flush()
        except urllib.error.HTTPError as e:
            try:
                body = e.read().decode("utf-8","replace")
                detail = json.loads(body).get("error",{}).get("message", body[:300])
            except: detail = str(e)
            try: self.wfile.write(("data: "+json.dumps({"error":f"Google API {e.code}: {detail}"})+"\n\n").encode()); self.wfile.flush()
            except: pass
        except Exception as e:
            try: self.wfile.write(("data: "+json.dumps({"error":"Google API: "+str(e)})+"\n\n").encode()); self.wfile.flush()
            except: pass

    def _stream_openai(self, model, system, messages):
        """Streaming via OpenAI API. Sleutel via instellingen of OPENAI_API_KEY env."""
        api_key = self.vault.get_api_key("openai")
        if not api_key:
            try: self.wfile.write(("data: "+json.dumps({"error":"OpenAI API-sleutel niet ingesteld — voeg toe via ⚙ Instellingen → API-sleutels, of: export OPENAI_API_KEY=sk-..."})+"\n\n").encode()); self.wfile.flush()
            except: pass
            return
        oai_msgs = [{"role":"system","content":system}] + \
                   [{"role":m["role"],"content":(m.get("content") or "").strip()}
                    for m in messages if m.get("role") in ("user","assistant") and (m.get("content") or "").strip()]
        if len(oai_msgs) < 2:
            oai_msgs.append({"role":"user","content":"(Vervolg)"})
        payload = {"model":model,"messages":oai_msgs,"stream":True,"max_tokens":4096}
        try:
            req = urllib.request.Request("https://api.openai.com/v1/chat/completions",
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type":"application/json","Authorization":"Bearer "+api_key},
                method="POST")
            with urllib.request.urlopen(req, timeout=120) as resp:
                for line in resp:
                    line = line.strip()
                    if not line: continue
                    try:
                        ls = line.decode("utf-8") if isinstance(line,bytes) else line
                        if ls.startswith("data: "):
                            data_s = ls[6:]
                            if data_s == "[DONE]":
                                self.wfile.write(("data: "+json.dumps({"delta":"","done":True})+"\n\n").encode("utf-8"))
                                self.wfile.flush(); break
                            ev = json.loads(data_s)
                            delta = ev.get("choices",[{}])[0].get("delta",{}).get("content","") or ""
                            self.wfile.write(("data: "+json.dumps({"delta":delta,"done":False})+"\n\n").encode("utf-8"))
                            self.wfile.flush()
                    except: pass
        except urllib.error.HTTPError as e:
            try:
                body = e.read().decode("utf-8","replace")
                detail = json.loads(body).get("error",{}).get("message", body[:300])
            except: detail = str(e)
            try: self.wfile.write(("data: "+json.dumps({"error":f"OpenAI {e.code}: {detail}"})+"\n\n").encode()); self.wfile.flush()
            except: pass
        except Exception as e:
            try: self.wfile.write(("data: "+json.dumps({"error":"OpenAI API: "+str(e)})+"\n\n").encode()); self.wfile.flush()
            except: pass

    def _stream_openrouter(self, model, system, messages):
        """Streaming via OpenRouter. Sleutel via instellingen of OPENROUTER_API_KEY env."""
        api_key = self.vault.get_api_key("openrouter")
        if not api_key:
            try: self.wfile.write(("data: "+json.dumps({"error":"OpenRouter API-sleutel niet ingesteld — voeg toe via ⚙ Instellingen → API-sleutels, of: export OPENROUTER_API_KEY=sk-or-..."})+"\n\n").encode()); self.wfile.flush()
            except: pass
            return
        or_msgs = [{"role":"system","content":system}] + \
                  [{"role":m["role"],"content":(m.get("content") or "").strip()}
                   for m in messages if m.get("role") in ("user","assistant") and (m.get("content") or "").strip()]
        if len(or_msgs) < 2:
            or_msgs.append({"role":"user","content":"(Vervolg)"})
        payload = {"model":model,"messages":or_msgs,"stream":True,"max_tokens":4096}
        try:
            req = urllib.request.Request("https://openrouter.ai/api/v1/chat/completions",
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type":"application/json",
                         "Authorization":"Bearer "+api_key,
                         "HTTP-Referer":"http://localhost:8899",
                         "X-Title":"Zettelkasten"},
                method="POST")
            with urllib.request.urlopen(req, timeout=120) as resp:
                for line in resp:
                    line = line.strip()
                    if not line: continue
                    try:
                        ls = line.decode("utf-8") if isinstance(line,bytes) else line
                        if ls.startswith("data: "):
                            data_s = ls[6:]
                            if data_s == "[DONE]":
                                self.wfile.write(("data: "+json.dumps({"delta":"","done":True})+"\n\n").encode("utf-8"))
                                self.wfile.flush(); break
                            ev = json.loads(data_s)
                            delta = ev.get("choices",[{}])[0].get("delta",{}).get("content","") or ""
                            self.wfile.write(("data: "+json.dumps({"delta":delta,"done":False})+"\n\n").encode("utf-8"))
                            self.wfile.flush()
                    except: pass
        except urllib.error.HTTPError as e:
            try:
                body = e.read().decode("utf-8","replace")
                detail = json.loads(body).get("error",{}).get("message", body[:300])
            except: detail = str(e)
            try: self.wfile.write(("data: "+json.dumps({"error":f"OpenRouter {e.code}: {detail}"})+"\n\n").encode()); self.wfile.flush()
            except: pass
        except Exception as e:
            try: self.wfile.write(("data: "+json.dumps({"error":"OpenRouter: "+str(e)})+"\n\n").encode()); self.wfile.flush()
            except: pass

    def _stream_mistral(self, model, system, messages):
        """Streaming via Mistral AI API (OpenAI-compatibel SSE formaat)."""
        api_key = self.vault.get_api_key("mistral")
        if not api_key:
            try: self.wfile.write(("data: "+json.dumps({"error":"Mistral API-sleutel niet ingesteld — voeg toe via ⚙ Instellingen → API-sleutels"})+"\n\n").encode()); self.wfile.flush()
            except: pass
            return
        msgs = [{"role":"system","content":system}] + \
               [{"role":m["role"],"content":(m.get("content") or "").strip()}
                for m in messages if m.get("role") in ("user","assistant") and (m.get("content") or "").strip()]
        if len(msgs) < 2:
            msgs.append({"role":"user","content":"(Vervolg)"})
        payload = {"model":model,"messages":msgs,"stream":True,"max_tokens":4096}
        try:
            req = urllib.request.Request("https://api.mistral.ai/v1/chat/completions",
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type":"application/json",
                         "Authorization":"Bearer "+api_key},
                method="POST")
            with urllib.request.urlopen(req, timeout=120) as resp:
                for line in resp:
                    line = line.strip()
                    if not line: continue
                    try:
                        ls = line.decode("utf-8") if isinstance(line,bytes) else line
                        if ls.startswith("data: "):
                            data_s = ls[6:]
                            if data_s == "[DONE]":
                                self.wfile.write(("data: "+json.dumps({"delta":"","done":True})+"\n\n").encode("utf-8"))
                                self.wfile.flush(); break
                            ev = json.loads(data_s)
                            delta = ev.get("choices",[{}])[0].get("delta",{}).get("content","") or ""
                            self.wfile.write(("data: "+json.dumps({"delta":delta,"done":False})+"\n\n").encode("utf-8"))
                            self.wfile.flush()
                    except: pass
        except urllib.error.HTTPError as e:
            try:
                body = e.read().decode("utf-8","replace")
                detail = json.loads(body).get("message", body[:300])
            except: detail = str(e)
            try: self.wfile.write(("data: "+json.dumps({"error":f"Mistral {e.code}: {detail}"})+"\n\n").encode()); self.wfile.flush()
            except: pass
        except Exception as e:
            try: self.wfile.write(("data: "+json.dumps({"error":"Mistral: "+str(e)})+"\n\n").encode()); self.wfile.flush()
            except: pass

    def _stream_ollama(self, model, system, messages):
        """Streaming via lokale Ollama."""
        payload={"model":model,"messages":[{"role":"system","content":system}]+messages,"stream":True}
        try:
            req=urllib.request.Request(self._ollama()+"/api/chat",
                data=json.dumps(payload).encode("utf-8"),
                headers={"Content-Type":"application/json"},method="POST")
            with urllib.request.urlopen(req,timeout=180) as resp:
                for line in resp:
                    line=line.strip()
                    if not line: continue
                    try:
                        c=json.loads(line)
                        d=c.get("message",{}).get("content","")
                        done=c.get("done",False)
                        self.wfile.write(("data: "+json.dumps({"delta":d,"done":done})+"\n\n").encode("utf-8"))
                        self.wfile.flush()
                        if done: break
                    except: pass
        except urllib.error.URLError as e:
            try: self.wfile.write(("data: "+json.dumps({"error":"Ollama niet bereikbaar: "+str(e)})+"\n\n").encode()); self.wfile.flush()
            except: pass
        except Exception as e:
            try: self.wfile.write(("data: "+json.dumps({"error":str(e)})+"\n\n").encode()); self.wfile.flush()
            except: pass

    def _call_llm_simple(self, model, prompt, max_tokens=600):
        """Roep een LLM aan met een simpele prompt, geef de response-tekst terug."""
        import urllib.request as _req2
        import urllib.error as _uerr
        response_text = ""

        if model.startswith("claude"):
            api_key = self.vault.get_api_key("anthropic")
            if api_key:
                payload = json.dumps({"model": model, "max_tokens": max_tokens,
                    "messages": [{"role":"user","content":prompt}]}).encode()
                req2 = _req2.Request("https://api.anthropic.com/v1/messages", data=payload,
                    headers={"Content-Type":"application/json","x-api-key":api_key,
                             "anthropic-version":"2023-06-01"}, method="POST")
                try:
                    with _req2.urlopen(req2, timeout=60) as r2:
                        d2 = json.loads(r2.read())
                    response_text = d2.get("content",[{}])[0].get("text","")
                except _uerr.HTTPError as e:
                    raise Exception(f"Anthropic API fout {e.code}: {e.read().decode()[:200]}")

        elif model.startswith("gpt") or model.startswith("o1") or model.startswith("o3") or model.startswith("o4"):
            api_key = self.vault.get_api_key("openai")
            if api_key:
                payload = json.dumps({"model": model, "max_tokens": max_tokens,
                    "messages": [{"role":"user","content":prompt}]}).encode()
                req2 = _req2.Request("https://api.openai.com/v1/chat/completions", data=payload,
                    headers={"Content-Type":"application/json","Authorization":"Bearer "+api_key},
                    method="POST")
                try:
                    with _req2.urlopen(req2, timeout=60) as r2:
                        d2 = json.loads(r2.read())
                    response_text = d2.get("choices",[{}])[0].get("message",{}).get("content","")
                except _uerr.HTTPError as e:
                    raise Exception(f"OpenAI API fout {e.code}: {e.read().decode()[:200]}")

        elif model.startswith("gemini"):
            api_key = self.vault.get_api_key("google")
            if api_key:
                payload = json.dumps({"contents":[{"role":"user","parts":[{"text":prompt}]}],
                    "generationConfig":{"maxOutputTokens":max_tokens}}).encode()
                url2 = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
                req2 = _req2.Request(url2, data=payload,
                    headers={"Content-Type":"application/json"}, method="POST")
                try:
                    with _req2.urlopen(req2, timeout=60) as r2:
                        d2 = json.loads(r2.read())
                    response_text = d2.get("candidates",[{}])[0].get("content",{}).get("parts",[{}])[0].get("text","")
                except _uerr.HTTPError as e:
                    raise Exception(f"Gemini API fout {e.code}: {e.read().decode()[:200]}")

        elif model.startswith("mistral") or "/" in model:  # Mistral direct of OpenRouter
            key_provider = "mistral" if model.startswith("mistral") else "openrouter"
            api_key = self.vault.get_api_key(key_provider)
            base_url = "https://api.mistral.ai/v1/chat/completions" if key_provider=="mistral" else "https://openrouter.ai/api/v1/chat/completions"
            if api_key:
                payload = json.dumps({"model": model, "max_tokens": max_tokens,
                    "messages": [{"role":"user","content":prompt}]}).encode()
                req2 = _req2.Request(base_url, data=payload,
                    headers={"Content-Type":"application/json","Authorization":"Bearer "+api_key},
                    method="POST")
                try:
                    with _req2.urlopen(req2, timeout=60) as r2:
                        d2 = json.loads(r2.read())
                    response_text = d2.get("choices",[{}])[0].get("message",{}).get("content","")
                except _uerr.HTTPError as e:
                    raise Exception(f"API fout {e.code}: {e.read().decode()[:200]}")

        # Fallback: Ollama lokaal
        if not response_text:
            r_ol = self._ollama_post("/api/generate",
                {"model": self._best_local_model(model), "prompt": prompt, "stream": False}, 120)
            response_text = r_ol.get("response","").strip()

        return response_text.strip()

    def _best_local_model(self, preferred=None):
        """Geeft het beste beschikbare Ollama-model terug.
        Negeert cloud-modellen (claude, gpt, gemini).
        Probeert preferred eerst, dan bekende goede modellen, dan eerste beschikbare."""
        CLOUD_PREFIXES = ("claude","gpt","gemini","mistral-api","openai")
        try:
            d = self._ollama_post("/api/tags", {}, 10)
            available = [m["name"].split(":")[0] for m in d.get("models", [])]
        except Exception:
            available = []

        # Goede volgorde van voorkeur voor tekst-taken
        PREFERRED_ORDER = [
            "llama3.2","llama3","llama3.1","llama2",
            "mistral","mixtral","phi3","phi","gemma2","gemma",
            "qwen2","qwen","deepseek","solar","vicuna","orca",
            "llama3.2-vision","llava",
        ]

        # Normaliseer preferred
        if preferred:
            pnorm = preferred.split(":")[0].lower()
            if not any(pnorm.startswith(c) for c in CLOUD_PREFIXES):
                # Preferred is een lokaal model — gebruik het direct
                return preferred

        # Zoek eerste match in voorkeursvolgorde
        avail_lower = [a.lower() for a in available]
        for want in PREFERRED_ORDER:
            for i, a in enumerate(avail_lower):
                if a == want or a.startswith(want):
                    return available[i]

        # Laatste redmiddel: eerste beschikbaar model dat geen cloud is
        for a in available:
            if not any(a.lower().startswith(c) for c in CLOUD_PREFIXES):
                return a

        return "llama3.2"  # ultieme fallback

    def _llm_summarize_pdf(self):
        body  = self._body()
        fname = body.get("filename","")
        model = body.get("model","")
        if not fname: return self._send(400,{"error":"filename vereist"})

        cfg          = self.vault.get_config()
        personal_use = cfg.get("pdf_personal_use", False)

        # ── Stap 1: tekst extraheren (personal_use wordt intern afgehandeld) ──
        text = self.vault.extract_pdf_text(fname, 40000) or ""
        print(f"[PDF-samenvatting] '{fname}': {len(text)} tekens geëxtraheerd (voor filter), "
              f"personal_use={personal_use}, model={model}", flush=True)

        # ── Stap 1b: verwijder DRM-watermerken en disclaimers uit de tekst ──────
        if text:
            import re as _re2

            # Stap A: verwijder e-mailadressen ZELF (niet de hele regel)
            text = _re2.sub(
                r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}',
                '', text
            )

            # Stap B: verwijder alleen expliciete disclaimer-zinnen
            # BELANGRIJK: patronen zijn zo specifiek mogelijk om inhoudstekst te sparen
            DISCLAIMER_PATTERNS = [
                r'(?i)[^\n]*(enkel|alleen|uitsluitend)\s+(bedoeld|bestemd)\s+voor[^\n]*\n?',
                r'(?i)[^\n]*bestemd\s+voor\s+persoonli[jk][^\n]*\n?',
                r'(?i)[^\n]*bedoeld\s+voor\s+persoonli[jk][^\n]*\n?',
                r'(?i)[^\n]*beperkt\s+tot\s+persoonli[jk][^\n]*\n?',
                r'(?i)[^\n]*voor\s+persoonli[jk]\s+gebruik\s+door[^\n]*\n?',
                r'(?i)[^\n]*geen\s+toegang\s+verleend[^\n]*\n?',
                r'(?i)[^\n]*anderen\s+dan\s+de\s+aangegeven[^\n]*\n?',
                r'(?i)[^\n]*niet\s+(verder\s+)?verspreid[^\n]*\n?',
                r'(?i)[^\n]*personal\s+use\s+only[^\n]*\n?',
                r'(?i)[^\n]*licensed?\s+(to|for|aan)\s+\w[^\n]*\n?',
            ]
            for pat in DISCLAIMER_PATTERNS:
                text = _re2.sub(pat, '\n', text)

            # Opschonen
            text = _re2.sub(r'[ \t]{2,}', ' ', text)
            text = _re2.sub(r'\n{3,}', '\n\n', text)
            text = text.strip()

        print(f"[PDF-samenvatting] '{fname}': {len(text)} tekens na disclaimer-filter", flush=True)

        # ── Stap 3: bepaal welk model gebruikt wordt ──────────────────────────
        # Online modellen worden direct gebruikt, lokaal via Ollama
        is_cloud = any(model.startswith(p) for p in
                       ("claude","gpt","gemini","o1","o3","o4",
                        "mistral","magistral","ministral")) or "/" in model

        def call_llm(prompt, images=None):
            """Roep het juiste model aan — cloud of lokaal. Gooit Exception bij fout."""
            import urllib.request as _r2
            import urllib.error as _e2

            def _http_error_msg(e):
                """Lees foutdetail uit HTTP-response body."""
                try:
                    body = e.read().decode("utf-8","replace")
                    detail = json.loads(body)
                    # Anthropic
                    if "error" in detail:
                        return detail["error"].get("message", body[:200])
                    # OpenAI / Gemini
                    if "message" in detail:
                        return detail["message"]
                    return body[:200]
                except Exception:
                    return str(e)

            if is_cloud and model.startswith("claude"):
                api_key = self.vault.get_api_key("anthropic")
                if not api_key:
                    raise Exception("Geen Anthropic API-sleutel ingesteld. Voeg deze toe via Instellingen → API-sleutels.")
                content = [{"type":"text","text":prompt}]
                if images:
                    for img_b64 in images[:3]:
                        content.insert(0,{"type":"image","source":{
                            "type":"base64","media_type":"image/png","data":img_b64}})
                payload = json.dumps({"model":model,"max_tokens":4000,
                                      "messages":[{"role":"user","content":content}]}).encode()
                req = _r2.Request("https://api.anthropic.com/v1/messages", data=payload,
                    headers={"Content-Type":"application/json","x-api-key":api_key,
                             "anthropic-version":"2023-06-01"}, method="POST")
                try:
                    with _r2.urlopen(req, timeout=90) as r2:
                        d2 = json.loads(r2.read())
                    result = d2.get("content",[{}])[0].get("text","")
                    if not result: raise Exception("Claude gaf een lege reactie terug.")
                    return result
                except _e2.HTTPError as e:
                    raise Exception(f"Claude API-fout {e.code}: {_http_error_msg(e)}")
                except _e2.URLError as e:
                    raise Exception(f"Claude niet bereikbaar: {e.reason}")

            elif is_cloud and (model.startswith("gpt") or model.startswith("o")):
                api_key = self.vault.get_api_key("openai")
                if not api_key:
                    raise Exception("Geen OpenAI API-sleutel ingesteld. Voeg deze toe via Instellingen → API-sleutels.")
                content = [{"type":"text","text":prompt}]
                if images:
                    for img_b64 in images[:3]:
                        content.insert(0,{"type":"image_url","image_url":{
                            "url":f"data:image/png;base64,{img_b64}"}})
                payload = json.dumps({"model":model,"max_tokens":4000,
                                      "messages":[{"role":"user","content":content}]}).encode()
                req = _r2.Request("https://api.openai.com/v1/chat/completions", data=payload,
                    headers={"Content-Type":"application/json","Authorization":"Bearer "+api_key},
                    method="POST")
                try:
                    with _r2.urlopen(req, timeout=90) as r2:
                        d2 = json.loads(r2.read())
                    result = d2.get("choices",[{}])[0].get("message",{}).get("content","")
                    if not result: raise Exception("OpenAI gaf een lege reactie terug.")
                    return result
                except _e2.HTTPError as e:
                    raise Exception(f"OpenAI API-fout {e.code}: {_http_error_msg(e)}")
                except _e2.URLError as e:
                    raise Exception(f"OpenAI niet bereikbaar: {e.reason}")

            elif is_cloud and model.startswith("gemini"):
                api_key = self.vault.get_api_key("google")
                if not api_key:
                    raise Exception("Geen Google API-sleutel ingesteld. Voeg deze toe via Instellingen → API-sleutels.")
                parts_list = [{"text":prompt}]
                if images:
                    for img_b64 in images[:3]:
                        parts_list.insert(0,{"inlineData":{"mimeType":"image/png","data":img_b64}})
                payload = json.dumps({"contents":[{"role":"user","parts":parts_list}],
                                      "generationConfig":{"maxOutputTokens":2000}}).encode()
                url2 = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
                req = _r2.Request(url2, data=payload,
                    headers={"Content-Type":"application/json"}, method="POST")
                try:
                    with _r2.urlopen(req, timeout=90) as r2:
                        d2 = json.loads(r2.read())
                    result = d2.get("candidates",[{}])[0].get("content",{}).get("parts",[{}])[0].get("text","")
                    if not result: raise Exception("Gemini gaf een lege reactie terug.")
                    return result
                except _e2.HTTPError as e:
                    raise Exception(f"Gemini API-fout {e.code}: {_http_error_msg(e)}")
                except _e2.URLError as e:
                    raise Exception(f"Gemini niet bereikbaar: {e.reason}")

            elif is_cloud and "/" in model:  # OpenRouter (Kimi, Llama, etc.)
                api_key = self.vault.get_api_key("openrouter")
                if not api_key:
                    raise Exception("Geen OpenRouter API-sleutel ingesteld. Voeg deze toe via Instellingen → API-sleutels.")
                payload = json.dumps({"model":model,"max_tokens":4000,
                                      "messages":[{"role":"user","content":prompt}]}).encode()
                req = _r2.Request("https://openrouter.ai/api/v1/chat/completions", data=payload,
                    headers={"Content-Type":"application/json","Authorization":"Bearer "+api_key,
                             "HTTP-Referer":"http://localhost:7842"}, method="POST")
                try:
                    with _r2.urlopen(req, timeout=90) as r2:
                        d2 = json.loads(r2.read())
                    result = d2.get("choices",[{}])[0].get("message",{}).get("content","")
                    if not result: raise Exception("OpenRouter gaf een lege reactie terug.")
                    return result
                except _e2.HTTPError as e:
                    raise Exception(f"OpenRouter API-fout {e.code}: {_http_error_msg(e)}")
                except _e2.URLError as e:
                    raise Exception(f"OpenRouter niet bereikbaar: {e.reason}")

            elif (model.startswith("mistral") or model.startswith("magistral")
                  or model.startswith("ministral")):
                api_key = self.vault.get_api_key("mistral")
                if not api_key:
                    raise Exception("Geen Mistral API-sleutel ingesteld. Voeg deze toe via Instellingen → API-sleutels.")
                payload = json.dumps({"model":model,"max_tokens":4000,
                                      "messages":[{"role":"user","content":prompt}]}).encode()
                req = _r2.Request("https://api.mistral.ai/v1/chat/completions", data=payload,
                    headers={"Content-Type":"application/json",
                             "Authorization":"Bearer "+api_key}, method="POST")
                try:
                    with _r2.urlopen(req, timeout=90) as r2:
                        d2 = json.loads(r2.read())
                    result = d2.get("choices",[{}])[0].get("message",{}).get("content","")
                    if not result: raise Exception("Mistral gaf een lege reactie terug.")
                    return result
                except _e2.HTTPError as e:
                    raise Exception(f"Mistral API-fout {e.code}: {_http_error_msg(e)}")
                except _e2.URLError as e:
                    raise Exception(f"Mistral niet bereikbaar: {e.reason}")

            else:
                # Lokaal Ollama
                local_model = self._best_local_model(model)
                try:
                    payload = {"model":local_model,"prompt":prompt,"stream":False}
                    if images:
                        payload["images"] = images[:2]
                    r = self._ollama_post("/api/generate", payload, 360)
                    result = r.get("response","").strip()
                    if not result: raise Exception(f"Ollama model '{local_model}' gaf een lege reactie.")
                    return result
                except Exception as e:
                    raise Exception(f"Ollama fout ({local_model}): {e}")

        # ── Stap 4: samenvatting maken ────────────────────────────────────────
        prompt = (
            "Je taak: schrijf een uitgebreide samenvatting van het DOCUMENT hieronder in Markdown.\n"
            "GEBRUIK ALLEEN de inhoudelijke informatie — verzin niets.\n"
            "Taal: gebruik de taal van het document (NL of EN).\n"
            "Lengte: MINIMAAL 400 en MAXIMAAL 900 woorden.\n\n"
            "STRIKT VERBODEN in de samenvatting — sla volledig over:\n"
            "- Elke tekst over 'persoonlijk gebruik', 'bestemd voor', 'bedoeld voor', 'beperkt tot'\n"
            "- Elke tekst over licenties, toegangsrechten, eigendom, vertrouwelijkheid\n"
            "- Zinnen die beginnen met 'Dit rapport/document is...'\n"
            "- E-mailadressen, namen van licentiehouders\n"
            "- Kopteksten, voetteksten, paginanummers, watermerken\n\n"
            "Structuur:\n"
            "## Inleiding\n"
            "Doel en context van het document (2-4 zinnen)\n\n"
            "## Kernpunten\n"
            "4-6 belangrijkste punten als bullets\n\n"
            "## Inhoud\n"
            "Samenvatting van de hoofdonderwerpen\n\n"
            "## Conclusies\n"
            "Conclusies en aanbevelingen\n\n"
            "## Vervolgvragen\n"
            "3-4 interessante vervolgvragen\n\n"
        )

        def call_llm_summarize(prompt_text, images=None):
            """Wrapper voor samenvatting — zelfde als call_llm maar duidelijk benoemd."""
            return call_llm(prompt_text, images)

        # Met tekst
        if text.strip():
            print(f"[PDF-samenvatting] Tekst beschikbaar ({len(text)} tekens), stuur naar model...", flush=True)
            try:
                full_prompt = (prompt +
                    "===BEGIN DOCUMENT===\n" + text[:35000] + "\n===EINDE DOCUMENT===\n\n"
                    "Schrijf nu de samenvatting op basis van bovenstaand document:")
                result = call_llm_summarize(full_prompt)
                if result:
                    print(f"[PDF-samenvatting] Samenvatting ontvangen ({len(result)} tekens)", flush=True)
                    return self._send(200,{"ok":True,"summary":result,"filename":fname})
            except Exception as e:
                print(f"[PDF-samenvatting] Fout: {e}", flush=True)
                return self._send(200,{"ok":False,"error":str(e),"summary":""})

        # Geen tekst → vision met meerdere pagina's
        images = self._pdf_pages_as_images(fname, max_pages=4)
        if images:
            vision_prompt = (
                "Dit zijn pagina's uit een PDF-document. "
                "Maak een Nederlandstalige samenvatting in Markdown van de inhoud. "
                "Gebruik ## voor hoofdsecties en - voor bullets. "
                "Wees zo concreet mogelijk over wat er in het document staat."
            )
            try:
                result = call_llm(vision_prompt, images=images)
                if result:
                    return self._send(200,{"ok":True,"summary":result,
                                           "filename":fname,"method":"vision"})
            except Exception as e:
                return self._send(200,{"ok":False,"error":str(e),"summary":""})

        msg = ("Geen tekst extraheerbaar uit dit PDF. "
               + ("Schakel 'Persoonlijk gebruik' in bij Instellingen → PDF om extra extractie-methoden te proberen."
                  if not personal_use else
                  "Dit PDF lijkt volledig gescand of zwaar beveiligd — vision-analyse mislukt ook."))
        return self._send(200,{"ok":False,"error":msg,"summary":""})

    def _pdf_pages_as_images(self, filename, max_pages=4):
        """Render PDF-pagina's naar base64 PNG's. Probeert pdftoppm, daarna pdf2image, daarna pypdf+PIL."""
        import base64, subprocess, tempfile, os
        pdf_path = self.vault.get_pdf_path(filename)
        if not pdf_path: return []
        images = []

        # Methode 1: pdftoppm (poppler)
        try:
            with tempfile.TemporaryDirectory() as td:
                out = os.path.join(td, "page")
                r = subprocess.run(
                    ["pdftoppm", "-r", "100", "-l", str(max_pages), "-png",
                     str(pdf_path), out],
                    capture_output=True, timeout=30)
                if r.returncode == 0:
                    imgs = sorted([f for f in os.listdir(td) if f.endswith(".png")])[:max_pages]
                    for img in imgs:
                        data = open(os.path.join(td, img), "rb").read()
                        images.append(base64.b64encode(data).decode())
                    if images: return images
        except Exception: pass

        # Methode 2: pdf2image (pip install pdf2image)
        try:
            from pdf2image import convert_from_path
            pages = convert_from_path(str(pdf_path), dpi=100,
                                      first_page=1, last_page=max_pages)
            import io
            for page in pages:
                buf = io.BytesIO()
                page.save(buf, format="PNG")
                images.append(base64.b64encode(buf.getvalue()).decode())
            if images: return images
        except Exception: pass

        # Methode 3: pypdf + PIL (alleen eerste pagina via annotatie-stream)
        try:
            import pypdf
            from PIL import Image
            import io
            reader = pypdf.PdfReader(str(pdf_path))
            for i, page in enumerate(reader.pages[:max_pages]):
                for name, obj in page.images:
                    img = Image.open(io.BytesIO(obj.data))
                    buf = io.BytesIO()
                    img.save(buf, format="PNG")
                    images.append(base64.b64encode(buf.getvalue()).decode())
                    break  # één afbeelding per pagina
            if images: return images
        except Exception: pass

        return []

    def _pdf_first_page_image(self, filename):
        """Legacy wrapper — geeft eerste pagina terug."""
        imgs = self._pdf_pages_as_images(filename, max_pages=1)
        return imgs[0] if imgs else None

    def _llm_summarize_note(self):
        """Genereer een AI-samenvatting van een notitie en geef die terug."""
        body  = self._body()
        nid   = body.get("note_id","")
        model = self._best_local_model(body.get("model",""))

        # Haal notitie op
        note = next((n for n in self.vault.load_notes() if n["id"]==nid), None)
        if not note:
            return self._send(404, {"error":"Notitie niet gevonden"})

        content = note.get("content","").strip()
        title   = note.get("title","")
        if len(content) < 50:
            return self._send(400, {"error":"Notitie te kort voor samenvatting"})

        prompt = (
            f"Schrijf een beknopte samenvatting van 4-6 zinnen van de onderstaande notitie.\n"
            f"Gebruik dezelfde taal als de notitie (NL of EN).\n"
            f"Beschrijf: het hoofdonderwerp, de kernpunten en de conclusie.\n"
            f"Begin de samenvatting NIET met 'Dit artikel' of 'Deze notitie' — begin direct met de inhoud.\n"
            f"Geef ALLEEN de samenvatting terug, geen extra tekst of uitleg.\n\n"
            f"Titel: {title}\n\n"
            f"Tekst:\n{content[:6000]}"
        )

        response_text = ""
        try:
            response_text = self._call_llm_simple(model, prompt, max_tokens=400)
        except Exception as e:
            return self._send(500, {"error": str(e)})

        summary = response_text.strip()
        if not summary:
            return self._send(500, {"error":"Model gaf geen samenvatting"})

        # ── Robuuste HTML/CSS sanitering ─────────────────────────────────────────
        import re as _re_s

        # Stap 1: volledige HTML-tags verwijderen (inclusief attributen)
        summary = _re_s.sub(r'<[^>]*>', '', summary)

        # Stap 2: CSS inline stijlen die overblijven na tag-verwijdering
        # bijv. "#8ac6f2;font-weight:bold;font-size:13px;margin-bottom:6px"
        summary = _re_s.sub(r'#?[0-9a-fA-F]{3,8};[\w-]+:[^;\n]{1,60}(?:;[\w-]+:[^;\n]{1,60})*', '', summary)
        summary = _re_s.sub(r'[\w-]+:[\w\s#.,%()]+;(?:[\w-]+:[\w\s#.,%()]+;?)*', '', summary)

        # Stap 3: HTML-entiteiten decoderen
        for enc, dec in [('&amp;','&'),('&lt;','<'),('&gt;','>'),('&nbsp;',' '),('&#39;',"'"),('&quot;','"')]:
            summary = summary.replace(enc, dec)

        # Stap 4: losse > of < tekens opruimen
        summary = _re_s.sub(r'(?<![\w\s])[<>](?![\w\s])', '', summary)

        # Stap 5: samenvatting-labels verwijderen die het model toevoegt
        summary = _re_s.sub(r'\*{0,2}(?:SAMENVATTING|Samenvatting|SUMMARY|Summary)\*{0,2}\s*[:\n]', '', summary)
        summary = _re_s.sub(r'={2,}\w+={2,}\s*', '', summary)

        # Stap 6: witruimte normaliseren
        summary = _re_s.sub(r'[ \t]{2,}', ' ', summary)
        summary = _re_s.sub(r'\n{3,}', '\n\n', summary)
        summary = summary.strip()
        return self._send(200, {"ok": True, "summary": summary})

    def _llm_link_reasons(self, query_content, candidates, model):
        """LLM beoordeelt kandidaat-notities en geeft reden per link."""
        cand_lines = []
        for i, c in enumerate(candidates[:8]):
            tags_str = ', '.join(c.get('tags', []) or [])
            title = c.get('title', '?')
            cand_lines.append(f'{i+1}. "{title}" (tags: {tags_str})')
        cand_text = "\n".join(cand_lines)
        preview = query_content[:1500]

        prompt = (
            "Je bent een Zettelkasten-assistent. Beoordeel of de onderstaande notities "
            "zinvol gelinkt kunnen worden aan de tekst.\n\n"
            "TEKST:\n" + preview + "\n\n"
            "KANDIDATEN:\n" + cand_text + "\n\n"
            "Geef per kandidaat een korte reden (max 8 woorden) waarom de link zinvol is, "
            "of \"nee\" als de link niet zinvol is.\n"
            "Antwoord ALLEEN als JSON-array in dezelfde volgorde, bijv:\n"
            '["bouwt voort op hetzelfde concept", "nee", "gedeelde methode"]'
        )

        response_text = self._call_llm_simple(model, prompt, max_tokens=400)
        if not response_text:
            return []

        import re as _re, json as _json
        clean = _re.sub(r'```json?|```', '', response_text).strip()
        try:
            raw = _json.loads(clean)
        except Exception:
            m = _re.search(r'\[.*?\]', clean, _re.DOTALL)
            if not m:
                return []
            try:
                raw = _json.loads(m.group())
            except Exception:
                return []

        result = []
        for i, cand in enumerate(candidates[:len(raw)]):
            reason = raw[i] if i < len(raw) else ""
            relevant = bool(reason and str(reason).strip().lower() not in
                           ("nee","no","geen","false",""))
            result.append({
                "id":       cand.get("id",""),
                "title":    cand.get("title",""),
                "reason":   reason if relevant else "",
                "relevant": relevant,
            })
        return result


    def _llm_describe_image(self):
        body=self._body()
        fname=body.get("filename",""); model=self._best_local_model(body.get("model",""))
        if not fname: return self._send(400,{"error":"filename vereist"})
        img=self.vault.image_as_base64(fname)
        if not img: return self._send(404,{"error":"Afbeelding niet gevonden"})
        prompt=("Beschrijf deze afbeelding in 3-5 zinnen Nederlands. "
                "Vermeld: wat er te zien is, kleuren, sfeer, eventuele tekst of symbolen.")
        try:
            r=self._ollama_post("/api/generate",{"model":model,"prompt":prompt,"images":[img["data"]],"stream":False},120)
            return self._send(200,{"ok":True,"description":r.get("response","").strip(),"filename":fname})
        except Exception as e:
            # Fallback naar llava
            try:
                r=self._ollama_post("/api/generate",{"model":"llava","prompt":prompt,"images":[img["data"]],"stream":False},120)
                return self._send(200,{"ok":True,"description":r.get("response","").strip(),"filename":fname})
            except Exception as e2:
                return self._send(200,{"ok":True,
                    "description":"Afbeelding: "+fname+" (vision model niet beschikbaar — ollama pull llama3.2-vision)",
                    "filename":fname,"warning":str(e2)})

    def _llm_mindmap(self):
        body=self._body()
        model=body.get("model","llama3.2-vision")
        ctx_notes=body.get("context_notes",[])
        ctx_pdfs=body.get("context_pdfs",[])
        parts=[]
        source_type="notes"  # "notes" | "pdf" | "mixed"

        if ctx_notes:
            for n in [x for x in self.vault.load_notes() if x["id"] in ctx_notes][:5]:
                parts.append("## "+n["title"]+"\n"+n["content"][:2000])
        if ctx_pdfs:
            source_type = "pdf" if not ctx_notes else "mixed"
            for pn in ctx_pdfs[:3]:
                txt=self.vault.extract_pdf_text(pn, 8000) or ""
                if txt.strip():
                    parts.append("=== DOCUMENT: "+pn+" ===\n"+txt)

        if not parts: return self._send(400,{"error":"Geen context"})

        if source_type in ("pdf","mixed"):
            prompt = (
                'Analyseer dit document grondig en maak een gedetailleerde mindmap die laat zien:\n'
                '1. Hoe het document is opgebouwd (structuur, hoofdstukken, secties)\n'
                '2. Wat de kernthema\'s en belangrijkste concepten zijn\n'
                '3. Hoe details, voorbeelden en argumenten zich vertalen in de boom\n'
                '4. Verbanden tussen onderdelen\n\n'
                'Geef ALLEEN geldige JSON terug (geen uitleg, geen backticks, geen markdown).\n'
                'Verplicht formaat met 3 niveaus:\n'
                '{\n'
                '  "root": "Documenttitel of kernthema",\n'
                '  "branches": [\n'
                '    {\n'
                '      "label": "Hoofdtak (sectie of thema)",\n'
                '      "color": "#hex",\n'
                '      "importance": "high|medium|low",\n'
                '      "children": [\n'
                '        {\n'
                '          "label": "Subtopic",\n'
                '          "details": ["detail1", "detail2"]\n'
                '        }\n'
                '      ]\n'
                '    }\n'
                '  ]\n'
                '}\n\n'
                'Gebruik 4-7 hoofdtakken. Elke tak heeft 2-5 subtopics. '
                'Elk subtopic heeft 1-3 details. '
                'Kies kleuren die de categorie weergeven: '
                '#8ac6f2 voor structuur/opbouw, #9fcf56 voor kernconcepten, '
                '#e5786d voor conclusies/bevindingen, #d4a4f7 voor methodes, '
                '#e8d44d voor voorbeelden, #f4bf75 voor aanbevelingen.\n\n'
                + "\n\n".join(parts)
            )
        else:
            prompt = (
                'Analyseer de notities en maak een mindmap van de kennisstructuur.\n'
                'Geef ALLEEN geldige JSON terug (geen uitleg, geen backticks).\n'
                'Formaat:\n'
                '{"root":"Overzicht","branches":[{"label":"Tak","color":"#hex",'
                '"importance":"high","children":[{"label":"Sub","details":["detail"]}]}]}\n'
                'Gebruik 3-6 takken met elk 2-4 subtopics.\n\n'
                + "\n\n".join(parts)
            )

        try:
            r=self._ollama_post("/api/generate",{"model":model,"prompt":prompt,"stream":False},180)
            raw=r.get("response","").strip()
            # Extracteer JSON — ook als model er tekst omheen zet
            m=re.search(r'\{[\s\S]*\}',raw)
            if m:
                mm=json.loads(m.group(0))
                # Normaliseer naar consistent formaat
                if "branches" not in mm: mm["branches"]=[]
                for b in mm["branches"]:
                    if "children" not in b: b["children"]=[]
                    for c in b["children"]:
                        if isinstance(c,str):
                            # Oud formaat: string kinderen
                            b["children"]=[{"label":ch,"details":[]} if isinstance(ch,str) else ch
                                           for ch in b["children"]]
                            break
                        if "details" not in c: c["details"]=[]
                return self._send(200,{"ok":True,"mindmap":mm,"source_type":source_type})
            return self._send(200,{"ok":False,"error":"Model gaf geen geldige JSON terug","raw":raw[:800]})
        except json.JSONDecodeError as e:
            return self._send(200,{"ok":False,"error":"JSON parse fout: "+str(e),"raw":raw[:400] if 'raw' in dir() else ""})
        except Exception as e:
            return self._send(200,{"ok":False,"error":str(e)})

    def _import_docx(self):
        """Converteer een geüpload Word-document naar Markdown en verwerk als URL-import."""
        import tempfile, os as _os, re as _re

        ct = self.headers.get("Content-Type","")
        if "multipart/form-data" not in ct:
            return self._send(400, {"error": "multipart/form-data vereist"})

        length = int(self.headers.get("Content-Length", 0))
        raw    = self.rfile.read(length)

        # Haal model op uit query-string
        from urllib.parse import parse_qs, urlparse as _up
        qs    = parse_qs(_up(self.path).query)
        model = qs.get("model", ["llama3.2-vision"])[0]

        # Zoek boundary (ondersteunt quoted en unquoted)
        bm = _re.search(r'boundary=([^\s;]+)', ct)
        if not bm:
            return self._send(400, {"error": "Geen boundary gevonden in Content-Type"})
        boundary = bm.group(1).strip('"').encode()

        # Splits op boundary
        parts = raw.split(b"--" + boundary)

        docx_data = None
        filename  = "document.docx"
        for part in parts:
            if b'filename=' not in part:
                continue
            # Bestandsnaam uit Content-Disposition
            m = _re.search(rb'filename=["\']?([^"\'\r\n]+)["\']?', part)
            if m:
                filename = m.group(1).decode("utf-8", "replace").strip()
            # Body staat na de dubbele CRLF
            sep = b"\r\n\r\n"
            if sep in part:
                body = part.split(sep, 1)[1]
                # Verwijder afsluitende CRLF en boundary-resten
                body = body.rstrip(b"\r\n-")
                docx_data = body
            break

        if not docx_data:
            return self._send(400, {"error": "Geen Word-bestand ontvangen"})

        # Schrijf tijdelijk naar schijf
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
            tf.write(docx_data)
            tmp_path = tf.name

        try:
            markdown, title = self._docx_to_markdown(tmp_path)
        finally:
            try: _os.unlink(tmp_path)
            except: pass

        if not markdown.strip():
            return self._send(400, {"error": "Kon geen tekst uit het Word-document halen"})

        # Samenvatting + opschoning via LLM (zelfde flow als URL-import)
        clean_md, summary = self._llm_clean_article(markdown, model)

        stem = _re.sub(r'\.(docx?|DOC[Xx]?)$', '', filename)
        if not title:
            title = stem

        return self._send(200, {
            "ok":      True,
            "title":   title,
            "md":      clean_md or markdown,
            "summary": summary,
            "source":  filename,
            "images":  [],
        })

    def _docx_to_markdown(self, path):
        """Converteer .docx naar Markdown. Geeft (markdown, title) terug."""
        import re as _re

        # ── Methode A: python-docx ────────────────────────────────────────────
        try:
            from docx import Document

            doc   = Document(path)
            lines = []
            title = ""

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    lines.append("")
                    continue
                try:
                    sname = para.style.name or ""
                except Exception:
                    sname = ""
                style = sname.lower()

                if style == "title" or style == "heading 1":
                    if not title: title = text
                    lines.append(f"# {text}")
                elif "heading 2" in style:
                    lines.append(f"## {text}")
                elif "heading 3" in style:
                    lines.append(f"### {text}")
                elif "heading" in style:
                    lines.append(f"#### {text}")
                elif "list" in style or sname.startswith("List"):
                    try:
                        indent = para.paragraph_format.left_indent
                        prefix = "  " * min(int((indent or 0) / 360000), 3)
                    except Exception:
                        prefix = ""
                    lines.append(f"{prefix}- {text}")
                elif style in ("quote", "intense quote", "block text"):
                    lines.append(f"> {text}")
                else:
                    lines.append(text)

            for table in doc.tables:
                if not table.rows: continue
                header = [c.text.strip() for c in table.rows[0].cells]
                lines.append("")
                lines.append("| " + " | ".join(header) + " |")
                lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                for row in table.rows[1:]:
                    lines.append("| " + " | ".join(c.text.strip() for c in row.cells) + " |")
                lines.append("")

            md = _re.sub(r'\n{3,}', '\n\n', "\n".join(lines)).strip()
            if not title and md:
                title = next((l.lstrip("#").strip() for l in md.splitlines() if l.strip()), "")[:80]
            return md, title

        except ImportError:
            print("[docx] python-docx niet gevonden — gebruik XML-fallback. "
                  "Installeer met: pip install python-docx", flush=True)

        # ── Methode B: ruwe XML uit .docx zip (geen externe packages) ────────
        try:
            import zipfile, xml.etree.ElementTree as ET

            NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            with zipfile.ZipFile(path, "r") as zf:
                xml_bytes = zf.read("word/document.xml")

            root  = ET.fromstring(xml_bytes)
            lines = []
            title = ""

            for para in root.iter(f"{{{NS}}}p"):
                # Stijlnaam
                pPr   = para.find(f".//{{{NS}}}pStyle")
                sname = pPr.get(f"{{{NS}}}val", "") if pPr is not None else ""
                style = sname.lower()

                # Alle tekst in de paragraaf
                text = "".join(t.text or "" for t in para.iter(f"{{{NS}}}t")).strip()
                if not text:
                    lines.append("")
                    continue

                if style in ("title", "heading1"):
                    if not title: title = text
                    lines.append(f"# {text}")
                elif style == "heading2":
                    lines.append(f"## {text}")
                elif style == "heading3":
                    lines.append(f"### {text}")
                elif style.startswith("heading"):
                    lines.append(f"#### {text}")
                elif "list" in style:
                    lines.append(f"- {text}")
                else:
                    lines.append(text)

            md = _re.sub(r'\n{3,}', '\n\n', "\n".join(lines)).strip()
            if not title and md:
                title = next((l.lstrip("#").strip() for l in md.splitlines() if l.strip()), "")[:80]
            print(f"[docx] XML-fallback: {len(md)} tekens geëxtraheerd", flush=True)
            return md, title

        except Exception as e:
            print(f"[docx] XML-fallback fout: {e}", flush=True)
            return "", ""

    def _llm_clean_article(self, text, model=""):
        """Maak een Markdown-samenvatting van een tekst via LLM.
        Geeft (opgeschoonde_tekst, samenvatting) terug.
        Bij lokale modellen of fouten: geeft (tekst, '') terug."""
        if not text.strip():
            return text, ""

        # Bepaal of we een cloud-model hebben
        is_cloud = any(model.startswith(p) for p in
                       ("claude","gpt","gemini","o1","o3","o4",
                        "mistral","magistral","ministral")) or "/" in model

        summary_prompt = (
            "Schrijf een beknopte samenvatting (3-5 zinnen) van onderstaande tekst. "
            "Gebruik de taal van de tekst (NL of EN). "
            "Geef ALLEEN de samenvatting terug, geen inleiding of labels.\n\n"
            f"TEKST:\n{text[:6000]}"
        )

        summary = ""
        try:
            if model.startswith("claude"):
                api_key = self.vault.get_api_key("anthropic")
                if api_key:
                    import urllib.request as _r2
                    payload = json.dumps({"model":model,"max_tokens":400,
                                          "messages":[{"role":"user","content":summary_prompt}]}).encode()
                    req = _r2.Request("https://api.anthropic.com/v1/messages", data=payload,
                        headers={"Content-Type":"application/json","x-api-key":api_key,
                                 "anthropic-version":"2023-06-01"}, method="POST")
                    with _r2.urlopen(req, timeout=30) as r:
                        summary = json.loads(r.read()).get("content",[{}])[0].get("text","")

            elif model.startswith("gpt") or model.startswith("o"):
                api_key = self.vault.get_api_key("openai")
                if api_key:
                    import urllib.request as _r2
                    payload = json.dumps({"model":model,"max_tokens":400,
                                          "messages":[{"role":"user","content":summary_prompt}]}).encode()
                    req = _r2.Request("https://api.openai.com/v1/chat/completions", data=payload,
                        headers={"Content-Type":"application/json","Authorization":"Bearer "+api_key},
                        method="POST")
                    with _r2.urlopen(req, timeout=30) as r:
                        summary = json.loads(r.read()).get("choices",[{}])[0].get("message",{}).get("content","")

            elif model.startswith("gemini"):
                api_key = self.vault.get_api_key("google")
                if api_key:
                    import urllib.request as _r2
                    payload = json.dumps({"contents":[{"role":"user","parts":[{"text":summary_prompt}]}],
                                          "generationConfig":{"maxOutputTokens":400}}).encode()
                    url2 = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
                    req = _r2.Request(url2, data=payload,
                        headers={"Content-Type":"application/json"}, method="POST")
                    with _r2.urlopen(req, timeout=30) as r:
                        summary = json.loads(r.read()).get("candidates",[{}])[0].get("content",{}).get("parts",[{}])[0].get("text","")

            elif model.startswith("mistral") or model.startswith("magistral"):
                api_key = self.vault.get_api_key("mistral")
                if api_key:
                    import urllib.request as _r2
                    payload = json.dumps({"model":model,"max_tokens":400,
                                          "messages":[{"role":"user","content":summary_prompt}]}).encode()
                    req = _r2.Request("https://api.mistral.ai/v1/chat/completions", data=payload,
                        headers={"Content-Type":"application/json","Authorization":"Bearer "+api_key},
                        method="POST")
                    with _r2.urlopen(req, timeout=30) as r:
                        summary = json.loads(r.read()).get("choices",[{}])[0].get("message",{}).get("content","")

            elif model.startswith("mistral") or "/" in model:  # Mistral direct of OpenRouter
                key_provider2 = "mistral" if model.startswith("mistral") else "openrouter"
                base_url2 = "https://api.mistral.ai/v1/chat/completions" if key_provider2=="mistral" else "https://openrouter.ai/api/v1/chat/completions"
                api_key = self.vault.get_api_key(key_provider2)
                if api_key:
                    import urllib.request as _r2
                    payload = json.dumps({"model":model,"max_tokens":400,
                                          "messages":[{"role":"user","content":summary_prompt}]}).encode()
                    req = _r2.Request("https://openrouter.ai/api/v1/chat/completions", data=payload,
                        headers={"Content-Type":"application/json","Authorization":"Bearer "+api_key,
                                 "HTTP-Referer":"http://localhost:7842"}, method="POST")
                    with _r2.urlopen(req, timeout=30) as r:
                        summary = json.loads(r.read()).get("choices",[{}])[0].get("message",{}).get("content","")

            else:
                # Lokaal Ollama
                try:
                    local = self._best_local_model(model)
                    r = self._ollama_post("/api/generate",
                        {"model":local,"prompt":summary_prompt,"stream":False}, 60)
                    summary = r.get("response","").strip()
                except Exception:
                    summary = ""

        except Exception as e:
            print(f"[docx-samenvatting] fout: {e}", flush=True)
            summary = ""

        return text, summary.strip()

    def _import_url(self):
        """Haal inhoud van een URL op en converteer naar Markdown — Instapaper-stijl."""
        body  = self._body()
        url   = body.get("url","").strip()
        model = body.get("model","llama3.2-vision")
        force = body.get("force", False)   # True = toch importeren ook al bestaat het
        if not url:
            return self._send(400,{"error":"url vereist"})

        # ── Server-side duplicate check ───────────────────────────────────────
        if not force:
            import re as _re_dup
            def _norm_url(u):
                try:
                    from urllib.parse import urlparse, urlencode, parse_qs, urlunparse
                    p = urlparse(u)
                    # Verwijder tracking-parameters
                    skip = {"utm_source","utm_medium","utm_campaign","utm_term",
                            "utm_content","fbclid","gclid","ref","source","si"}
                    qs = {k:v for k,v in parse_qs(p.query).items() if k not in skip}
                    clean = p._replace(query=urlencode(qs,doseq=True),
                                       path=p.path.rstrip("/"))
                    return urlunparse(clean).lower()
                except:
                    return u.rstrip("/").lower()

            target = _norm_url(url)
            for note in self.vault.load_notes():
                # Check 1: sourceUrl frontmatter
                if note.get("sourceUrl") and _norm_url(note["sourceUrl"]) == target:
                    return self._send(200, {
                        "ok": False, "duplicate": True,
                        "duplicate_id": note["id"],
                        "duplicate_title": note.get("title",""),
                        "error": f"Al geïmporteerd als: {note.get('title','onbekend')}"
                    })
                # Check 2: bron-link in content
                content = note.get("content","")
                urls_in_content = _re_dup.findall(r'\]\((https?://[^)]+)\)', content)
                if any(_norm_url(u) == target for u in urls_in_content):
                    return self._send(200, {
                        "ok": False, "duplicate": True,
                        "duplicate_id": note["id"],
                        "duplicate_title": note.get("title",""),
                        "error": f"Al geïmporteerd als: {note.get('title','onbekend')}"
                    })

        result = self._do_import_url(url, model)
        return self._send(200, result)

    def _do_import_url(self, url, model="llama3.2-vision"):
        """Kern-logica voor URL-import. Geeft dict terug (geen HTTP response).
        Downloadt ook afbeeldingen en slaat ze op in de vault/images map."""
        if not url.startswith(("http://","https://")):
            url = "https://" + url

        parsed_base = urlparse(url)
        base_url    = f"{parsed_base.scheme}://{parsed_base.netloc}"

        try:
            req = urllib.request.Request(url, headers={
                "User-Agent": "Mozilla/5.0 (compatible; Zettelkasten/1.0)",
                "Accept": "text/html,application/xhtml+xml",
                "Accept-Language": "nl,en;q=0.9",
            })
            with urllib.request.urlopen(req, timeout=15) as resp:
                raw_bytes = resp.read(500_000)
                charset   = "utf-8"
                ct = resp.headers.get("Content-Type","")
                if "charset=" in ct:
                    charset = ct.split("charset=")[-1].split(";")[0].strip()
                html = raw_bytes.decode(charset, errors="replace")
        except Exception as e:
            return {"ok":False,"error":"URL ophalen mislukt: "+str(e)}

        # ── HTML → tekst + afbeelding-URLs via stdlib html.parser ────────────
        from html.parser import HTMLParser

        class ArticleExtractor(HTMLParser):
            SKIP_TAGS    = {"script","style","nav","footer","header","aside",
                            "noscript","form","button","svg","iframe"}
            PARA_TAGS    = {"p","div","section","article","main",
                            "td","th","figcaption","blockquote","pre"}
            HEADING_TAGS = {"h1","h2","h3","h4","h5","h6"}
            LIST_TAGS    = {"li"}

            def __init__(self):
                super().__init__()
                self.skip_depth   = 0
                self.chunks       = []
                self.title        = ""
                self._in_title    = False
                self.img_srcs     = []
                self._img_seen    = set()
                self.og_image     = ""
                self._cur_heading = None

            @property
            def text_chunks(self): return self.chunks

            def _add_img(self, src):
                if not src or src.startswith("data:") or src in self._img_seen: return
                self._img_seen.add(src); self.img_srcs.append(src)

            def handle_starttag(self, tag, attrs):
                attr_d = dict(attrs)
                if tag == "title": self._in_title = True; return
                if tag in self.SKIP_TAGS: self.skip_depth += 1; return
                if self.skip_depth > 0: return

                if tag in self.HEADING_TAGS:
                    level = int(tag[1])
                    prefix = "#" * min(level + 1, 4)
                    self.chunks.append("")
                    self.chunks.append(prefix + " ")
                    self._cur_heading = tag; return

                if tag in self.LIST_TAGS:
                    self.chunks.append("- "); return

                if tag in self.PARA_TAGS:
                    if self.chunks and self.chunks[-1] not in ("",):
                        self.chunks.append("")
                    return

                if tag == "meta":
                    prop = attr_d.get("property","") or attr_d.get("name","")
                    if prop in ("og:image","og:image:secure_url","twitter:image"):
                        c = attr_d.get("content","").strip()
                        if c and not c.startswith("data:"): self.og_image = c

                if tag == "img":
                    for key in ("src","data-src","data-lazy-src","data-original",
                                "data-original-src","data-full-url","data-hi-res-src",
                                "data-src-medium","data-echo"):
                        c = attr_d.get(key,"")
                        if c and not c.startswith("data:"): self._add_img(c); break
                    srcset = attr_d.get("srcset","") or attr_d.get("data-srcset","")
                    if srcset:
                        parts = [p.strip().split()[0] for p in srcset.split(",") if p.strip()]
                        if parts: self._add_img(parts[-1])

                if tag == "source":
                    srcset = attr_d.get("srcset","") or attr_d.get("data-srcset","")
                    if srcset:
                        parts = [p.strip().split()[0] for p in srcset.split(",") if p.strip()]
                        if parts: self._add_img(parts[-1])

            def handle_endtag(self, tag):
                if tag == "title": self._in_title = False; return
                if tag in self.SKIP_TAGS and self.skip_depth > 0:
                    self.skip_depth -= 1; return
                if self.skip_depth > 0: return
                if tag in self.HEADING_TAGS:
                    self.chunks.append(""); self._cur_heading = None; return
                if tag in self.PARA_TAGS or tag in self.LIST_TAGS:
                    if self.chunks and self.chunks[-1] not in ("",):
                        self.chunks.append("")

            def handle_data(self, data):
                if self._in_title: self.title += data; return
                if self.skip_depth > 0: return
                t = data.strip()
                if not t: return
                # Plak tekst aan lopende heading of list-item
                if self.chunks and (self.chunks[-1].startswith("#") or self.chunks[-1] == "- "):
                    self.chunks[-1] += t
                else:
                    self.chunks.append(t)

        parser = ArticleExtractor()
        parser.feed(html)

        page_title = parser.title.strip() or url

        # Join met newlines — behoudt structuur
        raw_text = "\n".join(parser.chunks)
        raw_text = re.sub(r'\n{3,}', '\n\n', raw_text)   # max 2 opeenvolgende lege regels
        raw_text = re.sub(r'[ \t]+\n', '\n', raw_text)    # spaties voor newline
        # Zorg dat markdown-koppen altijd op een eigen regel staan
        raw_text = re.sub(r'([^\n])(#{1,4} )', r'\1\n\n\2', raw_text)
        # Lege regel na elke kop
        raw_text = re.sub(r'(#{1,4} [^\n]+)\n([^\n#\-])', r'\1\n\n\2', raw_text)
        raw_text = raw_text.strip()[:14000]

        if len(raw_text) < 100:
            return self._send(200,{"ok":False,
                "error":"Pagina bevat te weinig leesbare tekst (mogelijk JavaScript-only of betaalmuur)"})

        # ── Afbeeldingen downloaden ───────────────────────────────────────────
        IMAGE_MIME_TYPES = {"image/jpeg","image/png","image/gif","image/webp","image/svg+xml"}
        SKIP_PATTERNS    = ("logo","icon","avatar","badge","pixel","tracking","1x1",
                            "spacer","blank","transparent","spinner","loading")
        saved_images = []   # [{name, url_path}]

        # Voeg og:image toe als eerste (meest representatieve afbeelding van pagina)
        all_img_srcs = []
        if parser.og_image:
            all_img_srcs.append(parser.og_image)
        for s in parser.img_srcs:
            if s != parser.og_image:
                all_img_srcs.append(s)

        for raw_src in all_img_srcs[:40]:  # max 40 proberen (na filtering blijven er minder over)
            try:
                # Maak absolute URL
                if raw_src.startswith("//"):
                    img_url = parsed_base.scheme + ":" + raw_src
                elif raw_src.startswith("/"):
                    img_url = base_url + raw_src
                elif not raw_src.startswith("http"):
                    img_url = url.rsplit("/",1)[0] + "/" + raw_src
                else:
                    img_url = raw_src

                # Sla over als URL al een bekende tracker/icon is
                low = img_url.lower()
                if any(p in low for p in SKIP_PATTERNS):
                    continue

                # Download
                img_req = urllib.request.Request(img_url, headers={
                    "User-Agent":"Mozilla/5.0","Referer":url
                })
                with urllib.request.urlopen(img_req, timeout=10) as ir:
                    img_bytes = ir.read(5_000_000)  # max 5 MB
                    img_ct    = ir.headers.get("Content-Type","").split(";")[0].strip()

                # Filter op MIME
                if img_ct not in IMAGE_MIME_TYPES and not img_ct.startswith("image/"):
                    continue
                if len(img_bytes) < 800:    # kleiner dan 800 bytes = waarschijnlijk icon
                    continue

                # Bestandsnaam
                ext_map = {"image/jpeg":".jpg","image/png":".png","image/gif":".gif",
                           "image/webp":".webp","image/svg+xml":".svg"}
                ext = ext_map.get(img_ct, ".jpg")
                raw_name = img_url.split("?")[0].rsplit("/",1)[-1]
                if "." not in raw_name[-5:]:
                    raw_name = raw_name + ext
                # Prefix met domein voor uniciteit
                domain_prefix = parsed_base.netloc.replace("www.","").replace(".","_")
                fname = domain_prefix + "_" + raw_name[:60]

                saved = self.vault.save_image(fname, img_bytes)
                saved_images.append({"name": saved["name"], "url": saved["url"]})
            except Exception:
                continue  # sla mislukte downloads stilletjes over

        # ── Stap 0: ruwe tekst opschonen (regex, geen LLM nodig) ──────────────
        import re as _re

        # ── LinkedIn en andere login/paywall-fragmenten verwijderen ─────────────
        is_linkedin = "linkedin.com" in url

        LITERAL_JUNK = [
            # Cookie/privacy standaard
            "Lees meer in ons Cookiebeleid", "Lees meer in ons cookiebeleid",
            "meer in ons Cookiebeleid",
            "Lees meer in ons\nCookiebeleid", "Cookiebeleid . U kunt uw keuzen op elk gewenst moment bijwerken in uw instellingen",
            "Naar hoofdcontent gaan", "Naar hoofdinhoud gaan", "Ga naar hoofdinhoud",
            # LinkedIn UI-elementen
            "LinkedIn respecteert uw",
            "Selecteer Accepteren of Afwijzen om niet-essentiële",
            "U kunt uw keuzen op elk gewenst moment bijwerken in uw instellingen",
            "U kunt uw keuzen op elk gewenst moment bijwerken in uw\ninstellingen",
            "keuzen op elk gewenst moment bijwerken in uw instellingen",
            "Meld u aan om meer content weer te geven",
            "Maak uw gratis account of meld u aan om door te gaan met uw zoekopdracht",
            "Nog geen lid van LinkedIn? Word nu lid",
            "Door op Doorgaan te klikken om deel te nemen of u aan te melden",
            "gaat u akkoord met de gebruikersovereenkomst",
            "Meer informatie over onze privacybeleid",
            "Overgeslagen naar hoofdinhoud",
            "Topcommentaren", "Top comments",
            "Commentaren van mensen die u volgt",
            "Anderen bekeken ook", "Others also viewed",
            "Meer artikelen van", "More from",
            "Reacties weergeven", "View comments",
            "Vind ik leuk", "Like", "Reageren", "Delen", "Opslaan",
            "volgers", "verbindingen",
            "Volg", "Connect", "Bericht",
            "Uitgelicht door", "Featured by",
            "mensen vinden dit", "people find this",
            # Algemeen
            "Accept all cookies", "Alle cookies accepteren",
            "Manage preferences", "Voorkeuren beheren",
            "Cookie settings", "Cookie-instellingen",
            "Sign in to continue", "Log in to continue",
            "Subscribe to read", "Abonneer om te lezen",
            "This content is for subscribers",
        ]

        REGEX_JUNK = [
            r'(?i)lees meer in ons\s*\ncookiebeleid[^.]{0,300}[.\n]',
            r'(?i)cookiebeleid\s*\.\s*u kunt[^.]{0,200}[.\n]',
            r'(?i)u kunt uw keuzen[^.\n]{0,150}instellingen[^.\n]{0,50}[.\n]?',
            r'(?i)keuzen op elk gewenst moment[^.\n]{0,200}[.\n]?',
            r'(?i)naar hoofd(content|inhoud) gaan[^.]{0,100}[.\n]?',
            r'(?i)we (use|gebruiken) cookies?[^.]{0,250}[.\n]',
            r'(?i)cookie(s| settings| instellingen| beleid)[^.]{0,250}[.\n]',
            r'(?i)(accept|accepteer|decline|afwijzen|manage|beheer)\s+(all\s+)?(cookies?|preferences?)[^.]{0,200}[.\n]',
            r'(?i)by (clicking|continuing|using)[^.]{0,200}(terms|privacy|policy)[^.]{0,150}[.\n]',
            r'(?i)door (verder te gaan|gebruik te maken|te klikken)[^.]{0,200}[.\n]',
            r'(?i)gdpr[^.]{0,200}[.\n]',
            r'(?i)this site uses cookies[^.]{0,200}[.\n]',
            r'(?i)(subscribe|abonneer|sign up|aanmelden)[^.]{0,150}newsletter[^.]{0,100}[.\n]',
            r'(?i)\bpaywall\b[^.]{0,200}[.\n]',
            r'(?i)(register|registreer) (to|om) (read|lezen|access|verder)[^.]{0,200}[.\n]',
            # LinkedIn: getal + reacties/likes/shares
            r'\d+\s*(reacties?|commentaren?|comments?|likes?|shares?|reposts?)\b',
            # LinkedIn: "X volgers • Y verbindingen"
            r'\d[\d.,]*\s*(volgers?|followers?|verbindingen?|connections?)',
        ]

        clean_text = raw_text
        # Letterlijke fragmenten eerst (snelst)
        for junk in LITERAL_JUNK:
            # verwijder de zin inclusief alles tot het volgende punt of einde
            idx = clean_text.find(junk)
            while idx != -1:
                end = clean_text.find('.', idx + len(junk))
                if end == -1 or end - idx > 400:
                    end = idx + len(junk)
                clean_text = clean_text[:idx] + ' ' + clean_text[end+1:]
                idx = clean_text.find(junk)
        # Regex patronen
        for pat in REGEX_JUNK:
            clean_text = _re.sub(pat, ' ', clean_text)
        clean_text = _re.sub(r'\s{3,}', ' ', clean_text).strip()

        # ── Stap 1: één LLM-aanroep voor samenvatting + Markdown ─────────────
        linkedin_hint = (
            "\n\nLET OP — dit is een LinkedIn-artikel. Verwijder strikt:\n"
            "- Reactie-/like-/share-/repost-tellers (bijv. '42 reacties', '1.2K likes')\n"
            "- Profiel-blokken: naam, functie, bedrijf, volgers, verbindingen\n"
            "- Uitgelichte reacties, commentaren en antwoorden\n"
            "- 'Anderen bekeken ook', 'Meer van deze auteur', aanbevelingen\n"
            "- UI-labels: Volgen, Reageren, Delen, Opslaan, Bericht, Connect\n"
            "- Datum/tijdstempel van de post (niet van het artikel zelf)\n"
            "Behoud UITSLUITEND de doorlopende artikeltekst.\n"
        ) if is_linkedin else ""

        combined_prompt = (
            "Verwerk de onderstaande webpagina-tekst tot een nette notitie.\n\n"
            "STAP 1 — SAMENVATTING (verplicht, altijd):\n"
            "Schrijf 5-6 zinnen die het artikel samenvatten. Gebruik de taal van de tekst (NL of EN). "
            "Noem: het onderwerp, de kernboodschap en de conclusie. Geen koppen, geen bullets.\n\n"
            "STAP 2 — OPGESCHOONDE MARKDOWN:\n"
            "Zet de tekst om naar goed leesbare Markdown. Regels:\n"
            "- Gebruik ## voor hoofdsecties, ### voor subsecties\n"
            "- Gebruik **vet** voor kernbegrippen, *cursief* voor termen\n"
            "- Gebruik > voor citaten en quotes\n"
            "- Lege regel tussen elke alinea\n"
            "- Verwijder STRIKT: cookie-meldingen, navigatiemenu, login-verzoeken, \n"
            "  reclame, social-knoppen, reacties/comments, profielinfo, \n"
            "  'U kunt uw keuzen bijwerken', 'Lees meer in ons cookiebeleid'\n\n"
            "Geef je antwoord EXACT in dit format — geen extra tekst, geen uitleg:\n"
            "===SAMENVATTING===\n"
            "<5-6 zinnen samenvatting>\n"
            "===ARTIKEL===\n"
            "<opgeschoonde markdown>"
            + linkedin_hint + "\n\n"
            f"Titel: {page_title}\n\n"
            f"Tekst:\n{clean_text[:8000]}"
        )

        summary = ""
        md      = ""

        try:
            # Probeer eerst via het opgegeven model (kan online zijn)
            response_text = ""

            if model.startswith("claude"):
                api_key = self.vault.get_api_key("anthropic")
                if api_key:
                    import urllib.request as _req2
                    payload = json.dumps({
                        "model": model, "max_tokens": 3000,
                        "messages": [{"role":"user","content":combined_prompt}]
                    }).encode()
                    req2 = _req2.Request("https://api.anthropic.com/v1/messages",
                        data=payload,
                        headers={"Content-Type":"application/json",
                                 "x-api-key":api_key,
                                 "anthropic-version":"2023-06-01"},
                        method="POST")
                    with _req2.urlopen(req2, timeout=60) as r2:
                        d2 = json.loads(r2.read())
                    response_text = d2.get("content",[{}])[0].get("text","")

            elif model.startswith("gpt") or model.startswith("o1") or model.startswith("o3") or model.startswith("o4"):
                api_key = self.vault.get_api_key("openai")
                if api_key:
                    import urllib.request as _req2
                    payload = json.dumps({
                        "model": model, "max_tokens": 3000,
                        "messages": [{"role":"user","content":combined_prompt}]
                    }).encode()
                    req2 = _req2.Request("https://api.openai.com/v1/chat/completions",
                        data=payload,
                        headers={"Content-Type":"application/json",
                                 "Authorization":"Bearer "+api_key},
                        method="POST")
                    with _req2.urlopen(req2, timeout=60) as r2:
                        d2 = json.loads(r2.read())
                    response_text = d2.get("choices",[{}])[0].get("message",{}).get("content","")

            elif model.startswith("gemini"):
                api_key = self.vault.get_api_key("google")
                if api_key:
                    import urllib.request as _req2
                    payload = json.dumps({
                        "contents":[{"role":"user","parts":[{"text":combined_prompt}]}],
                        "generationConfig":{"maxOutputTokens":2000}
                    }).encode()
                    url2 = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
                    req2 = _req2.Request(url2,data=payload,
                        headers={"Content-Type":"application/json"},method="POST")
                    with _req2.urlopen(req2, timeout=60) as r2:
                        d2 = json.loads(r2.read())
                    response_text = d2.get("candidates",[{}])[0].get("content",{}).get("parts",[{}])[0].get("text","")

            elif model.startswith("mistral") or "/" in model:  # Mistral direct of OpenRouter
                key_provider2 = "mistral" if model.startswith("mistral") else "openrouter"
                base_url2 = "https://api.mistral.ai/v1/chat/completions" if key_provider2=="mistral" else "https://openrouter.ai/api/v1/chat/completions"
                api_key = self.vault.get_api_key(key_provider2)
                if api_key:
                    import urllib.request as _req2
                    payload = json.dumps({
                        "model": model, "max_tokens": 3000,
                        "messages": [{"role":"user","content":combined_prompt}]
                    }).encode()
                    req2 = _req2.Request(base_url2,
                        data=payload,
                        headers={"Content-Type":"application/json",
                                 "Authorization":"Bearer "+api_key},
                        method="POST")
                    with _req2.urlopen(req2, timeout=60) as r2:
                        d2 = json.loads(r2.read())
                    response_text = d2.get("choices",[{}])[0].get("message",{}).get("content","")

            # Fallback: Ollama lokaal
            if not response_text:
                r_ol = self._ollama_post("/api/generate",
                    {"model": self._best_local_model(model),
                     "prompt": combined_prompt, "stream": False}, 180)
                response_text = r_ol.get("response","").strip()

            # Parseer de response — zoek naar de markers, meerdere variaties
            response_text = response_text.strip()
            # Verwijder evt. markdown-fences
            response_text = _re.sub(r'^```\w*\n?', '', response_text).rstrip('`').strip()

            # Parseer markers — meerdere varianten ondersteunen
            def _extract(text):
                """Extraheer samenvatting + markdown uit model-response.
                Probeert 5 varianten zodat ook afwijkende model-outputs werken."""
                t = text.strip()

                # Stap 0: strip markdown code fences die sommige modellen toevoegen
                t = _re.sub(r'```(?:markdown|md|text)?\n?', '', t).strip()
                t = _re.sub(r'```\s*$', '', t, flags=_re.M).strip()

                # Variant 1: ===SAMENVATTING=== ... ===ARTIKEL===  (gewenst format)
                m1s = _re.search(r'={2,}\s*SAMENVATTING\s*={2,}\s*(.*?)\s*(?=={2,}\s*ARTIKEL)', t, _re.S|_re.I)
                m1a = _re.search(r'={2,}\s*ARTIKEL\s*={2,}\s*(.*?)(?:\s*={2,}|$)', t, _re.S|_re.I)
                if m1s and m1a:
                    s = m1s.group(1).strip()
                    a = m1a.group(1).strip()
                    if s and a: return s, a

                # Variant 2: **SAMENVATTING** / ## SAMENVATTING / SAMENVATTING:
                m2s = _re.search(
                    r'(?:\*{1,2}|#{1,4})?\s*(?:SAMENVATTING|SUMMARY)\s*(?:\*{1,2})?\s*[:=\-]*\s*\n+(.*?)'
                    r'(?=(?:\*{1,2}|#{1,4})?\s*(?:ARTIKEL|ARTICLE|INHOUD|CONTENT)\s*(?:\*{1,2})?\s*[:=\-]*)',
                    t, _re.S|_re.I)
                m2a = _re.search(
                    r'(?:\*{1,2}|#{1,4})?\s*(?:ARTIKEL|ARTICLE|INHOUD|CONTENT)\s*(?:\*{1,2})?\s*[:=\-]*\s*\n+(.*?)$',
                    t, _re.S|_re.I)
                if m2s:
                    s = m2s.group(1).strip()
                    a = m2a.group(1).strip() if m2a else ""
                    if s: return s, a

                # Variant 3: "Samenvatting:" gevolgd door tekst (zonder section break)
                m3 = _re.search(r'(?:Samenvatting|Summary):\s*([^\n]{40,}(?:\n(?![A-Z]{4,})[^\n]+){0,8})', t, _re.I)
                if m3:
                    s = m3.group(1).strip()
                    # Artikel is de rest
                    a = t[m3.end():].strip()
                    if s: return s, a or t

                # Variant 4: eerste alinea als samenvatting (als >=60 tekens, geen kop)
                paras = [p.strip() for p in t.split('\n\n') if p.strip()]
                if len(paras) >= 2:
                    first = paras[0]
                    # Sla over als het een kop of lijst is
                    if not _re.match(r'^#{1,4}\s|^-\s|^\*\s|^\d+\.\s', first) and len(first) > 60:
                        return first, '\n\n'.join(paras[1:])

                # Variant 5: geen scheiding gevonden - heel de tekst als artikel
                return "", t
            summary, md = _extract(response_text)

            if not md:
                md = clean_text[:6000]

            # ── Post-processing: markdown opschonen ──────────────────────────
            if md:
                # Zorg dat ## koppen altijd op een eigen regel staan
                md = _re.sub(r'([^\n])(#{1,4} )', r'\1\n\n\2', md)
                # Zorg dat een kop altijd een lege regel erna heeft
                md = _re.sub(r'(#{1,4} [^\n]+)\n([^\n#])', r'\1\n\n\2', md)
                # Verwijder dubbele lege regels (max 2)
                md = _re.sub(r'\n{3,}', '\n\n', md)
                # Verwijder CSS-rommel die modellen soms toevoegen
                md = _re.sub(r'#?[0-9a-fA-F]{0,8};?(?:[\w-]+:[^;">\n]{1,60};?){1,10}"?>', '', md)
                # Verwijder lege koppen (# alleen op een regel, overblijfsel van CSS-stripping)
                md = _re.sub(r'^#+\s*$', '', md, flags=_re.MULTILINE)
                # Opruimen: max 2 lege regels na stripping
                md = _re.sub(r'\n{3,}', '\n\n', md)
                md = md.strip()

            # ── Summary post-processing ──────────────────────────────────────
            if summary:
                # Strip markdown opmaak uit de samenvatting
                summary = _re.sub(r'#{1,4}\s+', '', summary)
                summary = _re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', summary)
                # CSS-rommel
                summary = _re.sub(r'#?[0-9a-fA-F]{0,8};?(?:[\w-]+:[^;">\n]{1,60};?){1,10}"?>', '', summary)
                summary = summary.strip()

        except Exception as e:
            md = clean_text[:6000]
            summary = ""

        return {
            "ok":       True,
            "title":    page_title,
            "url":      url,
            "summary":  summary,
            "markdown": md,
            "images":   saved_images,
        }

def get_local_ip():
    import socket
    try:
        s=socket.socket(socket.AF_INET,socket.SOCK_DGRAM); s.connect(("8.8.8.8",80))
        ip=s.getsockname()[0]; s.close(); return ip
    except: return "onbekend"


def main():
    parser=argparse.ArgumentParser(description="Zettelkasten VIM v4")
    parser.add_argument("--vault",      default=str(DEFAULT_VAULT))
    parser.add_argument("--port",       type=int,default=7842)
    parser.add_argument("--no-browser", action="store_true")
    parser.add_argument("--verbose",    action="store_true")
    parser.add_argument("--host",       default="0.0.0.0")
    parser.add_argument("--offline",    action="store_true",
                        help="Gebruik lokale vendor-bestanden i.p.v. CDN (geen internet vereist)")
    args=parser.parse_args()

    # Controleer vendor-bestanden bij --offline
    if args.offline:
        missing = [f for f in ["react.production.min.js","react-dom.production.min.js",
                                "pdf.min.js","pdf.worker.min.js"]
                   if not (VENDOR_DIR / f).exists()]
        if missing:
            print(f"\n⚠ Offline modus: ontbrekende vendor-bestanden in static/vendor/:")
            for m in missing: print(f"   - {m}")
            print(f"\n  Voer eerst uit:  cd static/vendor && bash download-vendors.sh\n")
            sys.exit(1)

    vault_path=Path(args.vault).expanduser().resolve()

    # ── Cloud-opslag detectie & waarschuwingen ─────────────────────────────
    vault_str = str(vault_path)
    cloud_type = None
    cloud_tips = []

    # OneDrive — macOS, Windows, Linux
    if any(x in vault_str for x in [
        "OneDrive", "onedrive",
        "Library/CloudStorage/OneDrive",
    ]):
        cloud_type = "OneDrive"
        cloud_tips = [
            "✓ Atomisch schrijven actief (veilig voor OneDrive-sync)",
            "✓ Bestanden zijn standaard Markdown — volledig compatibel",
            "⚠ Start de server NIET als OneDrive een grote sync uitvoert",
            "⚠ Stel OneDrive in op 'Always keep on this device' voor de vault-map",
        ]

    # Google Drive — macOS (Drive for Desktop), Windows (Drive letter G:/H:)
    elif any(x in vault_str for x in [
        "GoogleDrive", "Google Drive", "google-drive",
        "Library/CloudStorage/Google",
        "My Drive",
    ]) or (len(vault_str) == 3 and vault_str[1] == ":" and
           vault_str[0].upper() in "GHIJKLMNOPQRSTUVWXYZ"):
        cloud_type = "Google Drive"
        cloud_tips = [
            "✓ Atomisch schrijven actief (veilig voor Drive-sync)",
            "✓ Bestanden zijn standaard Markdown — volledig compatibel",
            "⚠ Schakel 'Mirror files' in (niet 'Stream files') voor betrouwbare toegang",
            "⚠ Op Windows: gebruik de stationsletter (bijv. G:\\Zettelkasten)",
        ]

    # iCloud Drive — macOS
    elif any(x in vault_str for x in [
        "iCloud Drive", "iCloudDrive",
        "Library/Mobile Documents/com~apple~CloudDocs",
    ]):
        cloud_type = "iCloud Drive"
        cloud_tips = [
            "✓ Atomisch schrijven actief",
            "⚠ iCloud kan bestanden 'verdampen' (evict) — zet optimalisatie UIT",
            "⚠ Ga naar Systeeminstellingen → iCloud → iCloud Drive → Optimaliseer Mac-opslag: UIT",
        ]

    # Dropbox
    elif any(x in vault_str for x in ["Dropbox", "dropbox"]):
        cloud_type = "Dropbox"
        cloud_tips = [
            "✓ Atomisch schrijven actief (veilig voor Dropbox-sync)",
            "✓ Dropbox ondersteunt .md bestanden zonder beperkingen",
            "⚠ Schakel 'Smart Sync' UIT voor de vault-map (anders offline niet beschikbaar)",
        ]

    ZKHandler.vault=VaultManager(vault_path)
    ZKHandler.verbose=args.verbose
    ZKHandler.offline=args.offline
    class ThreadingHTTPServer(ThreadingMixIn, HTTPServer):
        daemon_threads = True
        def handle_error(self, request, client_address):
            import sys
            exc = sys.exc_info()[1]
            if isinstance(exc, (BrokenPipeError, ConnectionResetError)):
                return  # client verbroken — geen stacktrace nodig
            super().handle_error(request, client_address)
    server=ThreadingHTTPServer((args.host,args.port),ZKHandler)
    local_ip=get_local_ip()
    offline_label = "JA (vendor/)" if args.offline else "nee (CDN)"
    cloud_block = ""
    if cloud_type:
        lines = "\n".join(f"  {t}" for t in cloud_tips)
        cloud_block = f"\n  Cloud   : {cloud_type}\n{lines}"

    print(f"""
╔══════════════════════════════════════════════════════╗
║        ZETTELKASTEN VIM  —  Python Server v4         ║
╚══════════════════════════════════════════════════════╝
  Vault   : {vault_path}
  Lokaal  : http://localhost:{args.port}
  Netwerk : http://{local_ip}:{args.port}
  Logging : {"aan" if args.verbose else "uit  (--verbose)"}{cloud_block}
  Offline : {offline_label}
  LLM     : ollama serve  +  ollama pull llama3.2-vision
  Stop    : Ctrl+C
""")
    if not args.no_browser:
        threading.Timer(0.8,lambda:webbrowser.open(f"http://localhost:{args.port}")).start()
    try: server.serve_forever()
    except KeyboardInterrupt: print("\nServer gestopt.")

if __name__=="__main__": main()
